#!/usr/bin/env python3
"""
Post-Separation Abuse & Court Case Analysis Tool
=================================================

Main entry point. Run this script to perform a full analysis on your case data.

Usage:
    python run_analysis.py --input ./input --state OH --county Franklin
    python run_analysis.py --input ./input/my_case.xlsx --gdrive-folder <FOLDER_URL>
    python run_analysis.py --interactive

Inputs accepted:
    - Excel workbooks (.xlsx/.xls)  — all sheets auto-classified
    - CSV files (.csv)
    - Text files (.txt)
    - Google Drive folder URL (requires credentials.json)

Outputs:
    - output/analysis_report.xlsx   — Full Excel report (new worksheet)
    - output/analysis_report.docx   — Formal Word document report
    - output/charts/                — Generated chart images
"""

import argparse
import json
import os
import sys
from datetime import datetime

import pandas as pd

from abuse_court_analyzer.data_loader import DataLoader
from abuse_court_analyzer.abuse_classifier import AbuseClassifier
from abuse_court_analyzer.cycle_predictor import CyclePredictor
from abuse_court_analyzer.court_analyzer import CourtAnalyzer
from abuse_court_analyzer.document_extractor import DocumentExtractor
from abuse_court_analyzer.document_compare import DocumentComparer
from abuse_court_analyzer.scorecards import OutcomeRatingSystem, OfficialScorecard
from abuse_court_analyzer.visualizations import ReportVisualizer
from abuse_court_analyzer.report_generator import ExcelReportGenerator, WordReportGenerator


def print_header():
    print("=" * 70)
    print("  POST-SEPARATION ABUSE & COURT CASE ANALYSIS TOOL  v1.0")
    print("  Forensic Pattern Analysis & Legally Defensible Reporting")
    print("=" * 70)
    print()


def _save_with_retry(report_obj, file_path, label="Report", max_retries=3):
    """Save a report with retry logic for Windows file-locking issues."""
    import time
    for attempt in range(max_retries + 1):
        try:
            report_obj.save()
            print(f"  {label} report: {file_path}")
            return
        except PermissionError:
            if attempt < max_retries:
                alt_path = file_path.replace(".", f"_{attempt + 1}.")
                print(f"  WARNING: {file_path} is locked. "
                      f"Retrying in 2s... (attempt {attempt + 1}/{max_retries})")
                time.sleep(2)
            else:
                # Save to alternate filename
                alt_path = file_path.replace(".", f"_backup.")
                print(f"  ERROR: {file_path} is locked by another program.")
                print(f"  Close Excel/Word and re-run, or find output at: {alt_path}")
                try:
                    report_obj.file_path = alt_path
                    report_obj.save()
                    print(f"  {label} report saved to: {alt_path}")
                except Exception as e2:
                    print(f"  Could not save {label} report: {e2}")


def run_analysis(input_path, state_code=None, county=None,
                 gdrive_folder=None, gdrive_links_col=None,
                 case_name="", case_number="",
                 compare_docs=None, output_dir="output"):
    """
    Execute the full analysis pipeline.

    Args:
        input_path: Path to input file or directory
        state_code: Two-letter state code (e.g., "OH")
        county: County name
        gdrive_folder: Google Drive folder URL/ID to scan
        gdrive_links_col: Column name in Excel containing Drive links
        case_name: Case name for the report
        case_number: Case number for the report
        compare_docs: Tuple of (doc_a_path, doc_b_path) for comparison
        output_dir: Output directory

    Returns:
        dict with paths to generated reports
    """
    os.makedirs(output_dir, exist_ok=True)
    os.makedirs(os.path.join(output_dir, "charts"), exist_ok=True)

    print_header()

    # ================================================================
    # 1. LOAD DATA
    # ================================================================
    print("[1/9] Loading data...")
    loader = DataLoader(input_path)

    if os.path.isdir(input_path):
        loader.load_directory(input_path)
    elif input_path.endswith((".xlsx", ".xls")):
        loader.load_excel(input_path)
    elif input_path.endswith(".csv"):
        loader.load_csv(input_path)
    elif input_path.endswith(".txt"):
        with open(input_path, "r") as f:
            loader.load_text(f.read())
    else:
        print(f"  Attempting to load as Excel: {input_path}")
        try:
            loader.load_excel(input_path)
        except Exception:
            loader.load_csv(input_path)

    classifications = loader.auto_classify_sheets()
    print(f"  Loaded {len(loader.raw_data)} data source(s):")
    for name, cls in classifications.items():
        df = loader.raw_data[name]
        print(f"    - {name}: {len(df)} rows, classified as '{cls}'")

    # ================================================================
    # 2. GOOGLE DRIVE INTEGRATION (if requested)
    # ================================================================
    gdrive_connector = None
    gdrive_documents = []

    if gdrive_folder or gdrive_links_col:
        print("\n[2/9] Connecting to Google Drive...")
        try:
            from abuse_court_analyzer.google_drive import GoogleDriveConnector
            gdrive_connector = GoogleDriveConnector()
            gdrive_connector.authenticate()
            print("  Authenticated successfully.")

            if gdrive_folder:
                print(f"  Scanning folder: {gdrive_folder}")
                gdrive_documents = gdrive_connector.gather_case_documents(gdrive_folder)
                print(f"  Found {len(gdrive_documents)} documents.")

            if gdrive_links_col and loader.abuse_data is not None:
                print(f"  Resolving Drive links in column '{gdrive_links_col}'...")
                loader.abuse_data = gdrive_connector.resolve_drive_links(
                    loader.abuse_data, link_column=gdrive_links_col
                )

        except ImportError:
            print("  Google Drive libraries not installed. Skipping.")
            print("  Install with: pip install google-auth google-auth-oauthlib google-api-python-client")
        except FileNotFoundError as e:
            print(f"  {e}")
            print("  Skipping Google Drive integration.")
        except Exception as e:
            print(f"  Google Drive error: {e}")
            print("  Continuing without Drive integration.")
    else:
        print("\n[2/9] Google Drive integration — skipped (no folder/links specified)")

    # ================================================================
    # 3. DOCUMENT EXTRACTION (from Drive or local)
    # ================================================================
    print("\n[3/9] Extracting facts from documents...")
    extractor = DocumentExtractor()

    if gdrive_documents:
        extractor.extract_from_documents(gdrive_documents)
        cross_doc = extractor.get_cross_document_analysis()
        print(f"  Extracted data from {cross_doc['total_documents']} documents")
        print(f"  Names found: {len(cross_doc['most_mentioned_names'])}")
        print(f"  Financial items: {cross_doc['financial_summary']['total_amounts_found']}")
    else:
        cross_doc = {}
        print("  No documents to extract from (provide --gdrive-folder or local docs)")

    # ================================================================
    # 4. CLASSIFY ABUSE INCIDENTS
    # ================================================================
    print("\n[4/9] Classifying abuse incidents...")
    classifier = AbuseClassifier()
    classified_df = None

    if loader.abuse_data is not None and not loader.abuse_data.empty:
        classified_df = classifier.classify_dataframe(loader.abuse_data)
        classified_df = classifier.calculate_casi(classified_df)
        formula_doc = classifier.generate_abuse_formula()
        print(f"  Classified {len(classified_df)} incidents")
        if "severity_level" in classified_df.columns:
            print(f"  Severity distribution:")
            for level, count in classified_df["severity_level"].value_counts().items():
                print(f"    {level}: {count}")
    else:
        formula_doc = classifier.generate_abuse_formula()
        print("  No abuse data found. Using formula documentation only.")
        # Try to use any available data
        if loader.raw_data:
            first_key = list(loader.raw_data.keys())[0]
            classified_df = classifier.classify_dataframe(loader.raw_data[first_key])
            if "severity_score" in classified_df.columns:
                classified_df = classifier.calculate_casi(classified_df)
                print(f"  Auto-classified {len(classified_df)} rows from '{first_key}'")

    # ================================================================
    # 5. PREDICT ABUSE CYCLES
    # ================================================================
    print("\n[5/9] Analyzing abuse cycles and generating predictions...")
    predictor = CyclePredictor()
    predictions = {}

    if classified_df is not None and not classified_df.empty:
        predictions = predictor.predict_next_cycle(classified_df)
        if "predictions" in predictions:
            print(f"  Generated {len(predictions['predictions'])} cycle predictions")
            for p in predictions["predictions"][:3]:
                print(f"    Cycle {p['cycle_number']}: ~{p['predicted_date']} "
                      f"(severity {p['predicted_severity']}, confidence {p['confidence']:.0%})")

        triggers = predictor.identify_triggers(classified_df)
        if triggers:
            print(f"  Identified {len(triggers)} potential trigger variables")
            predictions["triggers"] = triggers
    else:
        print("  Insufficient data for cycle prediction.")

    # ================================================================
    # 6. COURT CASE ANALYSIS
    # ================================================================
    print("\n[6/9] Analyzing court case patterns...")
    court = CourtAnalyzer(state_code=state_code, county=county)

    filings_df = loader.court_data if loader.court_data is not None else None
    if filings_df is None and classified_df is not None:
        filings_df = classified_df  # Use classified data as fallback

    duration = {}
    filing_patterns = {}
    violations = []
    costs = {}
    earnings = {}
    appeal_analysis = {}
    outcome_prediction = {}
    connections = []
    timeline = []

    # Enhanced analysis results
    analysis_365 = {}
    uccjea_analysis = {}
    best_interest = {}
    rule_11_violations = []
    rule_12_grounds = []
    rule_60_fraud = []
    judicial_code = {}
    per_person_scoring = {}
    dismissal_opportunities = []
    judge_rotation = {}
    conflict_of_interest = []
    detailed_timeline = []

    if filings_df is not None and not filings_df.empty:
        if loader.officials_data is not None:
            court.register_officials(loader.officials_data)
        court.track_official_activity(filings_df)
        officials_df = court.get_officials_summary()

        duration = court.analyze_case_duration(filings_df)
        timeline = court.build_timeline(filings_df)
        filing_patterns = court.analyze_filing_patterns(filings_df)
        violations = court.detect_procedural_violations(filings_df)
        costs = court.analyze_costs(loader.financial_data, filings_df)
        earnings = court.calculate_official_earnings()
        appeal_analysis = court.analyze_appeals(filings_df=filings_df)
        outcome_prediction = court.predict_outcomes(filings_df)
        connections = court.analyze_connections(loader.officials_data)

        # Enhanced SC-specific analysis
        if hasattr(court, 'analyze_365_day_order'):
            analysis_365 = court.analyze_365_day_order(filings_df)
            print(f"  365-day analysis: {analysis_365.get('status', 'N/A')}")

        if hasattr(court, 'analyze_uccjea_jurisdiction'):
            uccjea_analysis = court.analyze_uccjea_jurisdiction(
                addresses_data=loader.addresses_data, filings_df=filings_df)
            print(f"  UCCJEA issues: {len(uccjea_analysis.get('issues', []))}")

        if hasattr(court, 'analyze_best_interest_factors'):
            best_interest = court.analyze_best_interest_factors(
                classified_abuse_df=classified_df, filings_df=filings_df,
                officials_df=loader.officials_data)
            print(f"  Best interest factors analyzed: {len(best_interest.get('factors', []))}")

        if hasattr(court, 'detect_rule_11_violations'):
            rule_11_violations = court.detect_rule_11_violations(
                filings_df, detailed_hearings_df=loader.detailed_hearings_data,
                violations_df=loader.violations_data)
            print(f"  Rule 11 violations: {len(rule_11_violations)}")

        if hasattr(court, 'detect_rule_12_dismissal_grounds'):
            rule_12_grounds = court.detect_rule_12_dismissal_grounds(
                filings_df, addresses_df=loader.addresses_data,
                detailed_hearings_df=loader.detailed_hearings_data)
            print(f"  Rule 12 dismissal grounds: {len(rule_12_grounds)}")

        if hasattr(court, 'detect_rule_60_fraud'):
            rule_60_fraud = court.detect_rule_60_fraud(filings_df)
            print(f"  Rule 60 fraud indicators: {len(rule_60_fraud)}")

        if hasattr(court, 'analyze_judicial_code_compliance'):
            judicial_code = court.analyze_judicial_code_compliance(filings_df)
            print(f"  Judicial code issues: {len(judicial_code.get('violations', []))}")

        if hasattr(court, 'analyze_per_person_abuse_scoring') and classified_df is not None:
            per_person_scoring = court.analyze_per_person_abuse_scoring(classified_df)
            print(f"  Persons scored: {len(per_person_scoring)}")

        if hasattr(court, 'identify_dismissal_opportunities'):
            dismissal_opportunities = court.identify_dismissal_opportunities(
                filings_df, addresses_data=loader.addresses_data)
            print(f"  Dismissal opportunities: {len(dismissal_opportunities)}")

        if hasattr(court, 'analyze_judge_rotation'):
            judge_rotation = court.analyze_judge_rotation(filings_df)

        if hasattr(court, 'detect_conflict_of_interest'):
            conflict_of_interest = court.detect_conflict_of_interest(loader.officials_data)

        if hasattr(court, 'generate_detailed_event_timeline'):
            detailed_timeline = court.generate_detailed_event_timeline(
                filings_df, classified_df=classified_df,
                violations_df=loader.violations_data)
            print(f"  Detailed timeline events: {len(detailed_timeline)}")

        print(f"  Case duration: {duration.get('duration_years', 'N/A')} years")
        print(f"  Total filings: {duration.get('total_filings', 'N/A')}")
        print(f"  Officials tracked: {len(court.officials_registry)}")
        print(f"  Violations found: {len(violations)}")
        print(f"  Connections found: {len(connections)}")
    else:
        officials_df = pd.DataFrame()
        print("  No court filing data found.")

    # ================================================================
    # 7. SCORECARDS & RATINGS
    # ================================================================
    print("\n[7/9] Generating scorecards and outcome ratings...")

    # Outcome rating
    rating_system = OutcomeRatingSystem()
    case_score_data = {
        "case_duration": max(0, 100 - duration.get("duration_years", 0) * 15) if duration else 50,
        "filing_volume": max(0, 100 - duration.get("total_filings", 0) * 1) if duration else 50,
        "judge_history": 50,  # Neutral default
        "attorney_experience": 50,
        "gal_involvement": 50,
        "abuse_severity": max(0, 100 - (classified_df["severity_score"].mean() * 10 if classified_df is not None and "severity_score" in classified_df.columns else 50)),
        "financial_disparity": 50,
        "children_involved": 50,
        "violation_count": max(0, 100 - len(violations) * 10),
    }
    case_rating = rating_system.score_case_outcome(case_score_data)
    print(f"  Case outcome score: {case_rating['overall_score']}/100 ({case_rating['risk_level']} risk)")

    # Official scorecards
    scorecard_gen = OfficialScorecard()
    if court.officials_registry:
        scorecards = scorecard_gen.generate_all_scorecards(
            court.officials_registry, earnings
        )
        scorecards_df = scorecard_gen.get_scorecards_dataframe()
        print(f"  Generated {len(scorecards)} official scorecards")
    else:
        scorecards_df = pd.DataFrame()

    # ================================================================
    # 7b. DOCUMENT COMPARISON (if requested)
    # ================================================================
    comparison_result = None
    if compare_docs and len(compare_docs) == 2:
        print("\n  Comparing documents side-by-side...")
        comparer = DocumentComparer()
        doc_a_text = _read_doc_text(compare_docs[0], gdrive_connector)
        doc_b_text = _read_doc_text(compare_docs[1], gdrive_connector)
        comparison_result = comparer.compare(
            doc_a_text, doc_b_text,
            label_a=os.path.basename(str(compare_docs[0])),
            label_b=os.path.basename(str(compare_docs[1]))
        )
        print(f"  Similarity: {comparison_result['similarity_score']:.2%}")
        print(f"  Contradictions found: {len(comparison_result.get('contradictions', []))}")

    # ================================================================
    # 8. GENERATE VISUALIZATIONS
    # ================================================================
    print("\n[8/9] Generating charts and visualizations...")
    viz = ReportVisualizer(output_dir=os.path.join(output_dir, "charts"))

    if classified_df is not None and not classified_df.empty:
        viz.plot_abuse_timeline(classified_df)
        viz.plot_abuse_type_distribution(classified_df)
        viz.plot_severity_heatmap(classified_df)

    if predictions and "predictions" in predictions:
        viz.plot_cycle_prediction(predictions["predictions"], classified_df)

    if filings_df is not None and not filings_df.empty:
        viz.plot_filing_timeline(filings_df)

    if not officials_df.empty:
        viz.plot_officials_involvement(officials_df)

    if costs:
        viz.plot_financial_breakdown(costs)

    if violations:
        viz.plot_violation_summary(violations)

    if not scorecards_df.empty:
        viz.plot_scorecards(scorecards_df)

    if connections:
        viz.plot_connection_network(connections)

    if comparison_result:
        viz.plot_document_comparison(
            comparison_result.get("side_by_side", []),
            comparison_result.get("document_a", "Doc A"),
            comparison_result.get("document_b", "Doc B"),
        )

    print(f"  Generated {len(viz.generated_charts)} charts")

    # ================================================================
    # 9. GENERATE REPORTS
    # ================================================================
    print("\n[9/9] Generating reports...")

    # Executive summary
    exec_summary = {
        "case_name": case_name or "Not specified",
        "case_number": case_number or "Not specified",
        "jurisdiction": f"{county or 'N/A'}, {state_code or 'N/A'}",
        "report_date": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "data_sources": len(loader.raw_data),
        "total_incidents_analyzed": len(classified_df) if classified_df is not None else 0,
        "case_duration": duration,
        "abuse_severity_summary": {
            "avg_severity": round(classified_df["severity_score"].mean(), 2) if classified_df is not None and "severity_score" in classified_df.columns else "N/A",
            "max_severity": round(classified_df["severity_score"].max(), 2) if classified_df is not None and "severity_score" in classified_df.columns else "N/A",
            "escalation_trend": predictions.get("analysis_summary", {}).get("escalation", {}).get("trend", "N/A"),
        },
        "case_outcome_score": f"{case_rating['overall_score']}/100 ({case_rating['risk_level']} risk)",
        "violations_found": len(violations),
        "officials_tracked": len(court.officials_registry),
        "connections_found": len(connections),
    }

    # --- Excel Report ---
    xlsx_path = os.path.join(output_dir, "analysis_report.xlsx")
    excel = ExcelReportGenerator(xlsx_path)
    excel.add_executive_summary(exec_summary)

    if classified_df is not None:
        excel.add_abuse_analysis(classified_df)

    if predictions:
        excel.add_cycle_predictions(predictions)

    excel.add_formula_documentation(formula_doc)

    if timeline:
        excel.add_court_timeline(timeline)

    if not officials_df.empty:
        excel.add_officials_analysis(officials_df, earnings)

    if violations:
        excel.add_violations(violations)

    if costs:
        excel.add_financial_analysis(costs)

    if connections:
        excel.add_connections(connections)

    if not scorecards_df.empty:
        excel.add_scorecards(scorecards_df)

    if comparison_result:
        excel.add_document_comparison(comparison_result)

    excel.add_case_outcome_prediction({
        "case_rating": case_rating,
        "court_projection": outcome_prediction,
    })

    # Enhanced analysis sheets
    if hasattr(excel, 'add_detailed_event_timeline') and detailed_timeline:
        excel.add_detailed_event_timeline(detailed_timeline)

    if hasattr(excel, 'add_365_day_analysis') and analysis_365:
        excel.add_365_day_analysis(analysis_365)

    if hasattr(excel, 'add_uccjea_analysis') and uccjea_analysis:
        excel.add_uccjea_analysis(uccjea_analysis)

    if hasattr(excel, 'add_best_interest_factors') and best_interest:
        excel.add_best_interest_factors(best_interest)

    if hasattr(excel, 'add_rule_11_60_analysis') and (rule_11_violations or rule_60_fraud):
        excel.add_rule_11_60_analysis(rule_11_violations, rule_60_fraud)

    # Rule 12 dismissal grounds sheet
    if rule_12_grounds:
        excel.add_raw_data_sheet("Rule_12_Grounds",
                                 pd.DataFrame(rule_12_grounds))

    if hasattr(excel, 'add_judicial_code_analysis') and judicial_code:
        # Convert violations dict to the list format expected by report generator
        jc_violations = judicial_code.get("violations", [])
        if jc_violations:
            excel.add_raw_data_sheet("Judicial_Code_Violations", pd.DataFrame(jc_violations))

    if hasattr(excel, 'add_per_person_scoring') and per_person_scoring:
        excel.add_per_person_scoring(per_person_scoring)

    if hasattr(excel, 'add_dismissal_opportunities') and dismissal_opportunities:
        excel.add_dismissal_opportunities(dismissal_opportunities)

    if hasattr(excel, 'add_judge_rotation_analysis') and judge_rotation:
        excel.add_judge_rotation_analysis(judge_rotation)

    if hasattr(excel, 'add_conflict_of_interest') and conflict_of_interest:
        excel.add_conflict_of_interest(conflict_of_interest)

    # Add raw data sheets for querying
    if loader.visitation_data is not None:
        excel.add_raw_data_sheet("Visitation_Tracking", loader.visitation_data)
    if loader.addresses_data is not None:
        excel.add_raw_data_sheet("Addresses_Jurisdiction", loader.addresses_data)
    if loader.case_numbers_data is not None:
        excel.add_raw_data_sheet("Case_Numbers", loader.case_numbers_data)
    if loader.comprehensive_timeline_data is not None:
        excel.add_raw_data_sheet("Full_Timeline", loader.comprehensive_timeline_data)
    if loader.detailed_hearings_data is not None:
        excel.add_raw_data_sheet("Detailed_Hearings", loader.detailed_hearings_data)
    if loader.violations_data is not None:
        excel.add_raw_data_sheet("Violations_Detail", loader.violations_data)
    if loader.connections_data is not None:
        excel.add_raw_data_sheet("Connections", loader.connections_data)
    if hasattr(loader, 'attorney_consultations_data') and loader.attorney_consultations_data is not None:
        excel.add_raw_data_sheet("Attorney_Consultations", loader.attorney_consultations_data)
    if hasattr(loader, 'uccjea_violations_data') and loader.uccjea_violations_data is not None:
        excel.add_raw_data_sheet("UCCJEA_Violations", loader.uccjea_violations_data)
    if hasattr(loader, 'common_pleas_data') and loader.common_pleas_data is not None:
        excel.add_raw_data_sheet("Common_Pleas", loader.common_pleas_data)
    if hasattr(loader, 'visitation_costs_data') and loader.visitation_costs_data is not None:
        excel.add_raw_data_sheet("Visitation_Costs", loader.visitation_costs_data)

    # Rule 11 violations sheet
    if rule_11_violations:
        excel.add_raw_data_sheet("Rule_11_Violations", pd.DataFrame(rule_11_violations))

    # Dismissal opportunities summary
    if dismissal_opportunities:
        excel.add_raw_data_sheet("Dismissal_Opportunities", pd.DataFrame(dismissal_opportunities))

    excel.add_methodology()

    if viz.generated_charts:
        excel.add_charts_sheet(viz.generated_charts)

    _save_with_retry(excel, xlsx_path, "Excel")

    # --- Word Report ---
    docx_path = os.path.join(output_dir, "analysis_report.docx")
    word = WordReportGenerator(docx_path)
    word.add_title_page(case_name, case_number, state_code or "", county or "")
    word.add_table_of_contents()
    word.add_executive_summary(exec_summary)

    word.add_abuse_analysis_section(classified_df, formula_doc)

    if predictions:
        word.add_cycle_prediction_section(predictions)

    word.add_court_analysis_section(duration, filing_patterns, timeline)
    word.add_officials_section(officials_df, earnings, scorecards_df)
    word.add_violations_section(violations)
    word.add_financial_section(costs)
    word.add_connections_section(connections)
    word.add_outcome_prediction_section({
        "case_rating": case_rating,
        "court_projection": outcome_prediction,
    })

    if comparison_result:
        word.add_document_comparison_section(comparison_result)

    if viz.generated_charts:
        word.add_charts(viz.generated_charts)

    word.add_methodology_section()
    _save_with_retry(word, docx_path, "Word")

    print("\n" + "=" * 70)
    print("  ANALYSIS COMPLETE")
    print(f"  Reports saved to: {os.path.abspath(output_dir)}/")
    print("=" * 70)

    return {
        "excel_report": xlsx_path,
        "word_report": docx_path,
        "charts": viz.generated_charts,
        "executive_summary": exec_summary,
    }


def _read_doc_text(path_or_url, gdrive_connector=None):
    """Read document text from a local path or Google Drive URL."""
    path = str(path_or_url)

    if "drive.google" in path and gdrive_connector:
        result = gdrive_connector.read_file_text(path)
        return result

    if os.path.exists(path):
        ext = os.path.splitext(path)[1].lower()
        if ext == ".txt":
            with open(path, "r") as f:
                return {"title": os.path.basename(path), "text": f.read()}
        elif ext in (".docx",):
            from docx import Document
            doc = Document(path)
            text = "\n".join(p.text for p in doc.paragraphs)
            return {"title": os.path.basename(path), "text": text}
        elif ext == ".pdf":
            try:
                import fitz
                doc = fitz.open(path)
                text = ""
                for page in doc:
                    text += page.get_text()
                doc.close()
                return {"title": os.path.basename(path), "text": text}
            except ImportError:
                return {"title": os.path.basename(path), "text": "[PDF — install pymupdf]"}
        else:
            with open(path, "r", errors="ignore") as f:
                return {"title": os.path.basename(path), "text": f.read()}

    return {"title": str(path), "text": ""}


def main():
    parser = argparse.ArgumentParser(
        description="Post-Separation Abuse & Court Case Analysis Tool",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python run_analysis.py --input ./input --state OH --county Franklin
  python run_analysis.py --input case_data.xlsx --gdrive-folder "https://drive.google.com/drive/folders/..."
  python run_analysis.py --input ./input --compare doc1.txt doc2.txt
  python run_analysis.py --interactive
        """
    )

    parser.add_argument("--input", "-i", default="./input",
                        help="Path to input file or directory (default: ./input)")
    parser.add_argument("--state", "-s", default=None,
                        help="Two-letter state code (e.g., OH, TX, CA)")
    parser.add_argument("--county", "-c", default=None,
                        help="County name")
    parser.add_argument("--case-name", default="",
                        help="Case name for the report")
    parser.add_argument("--case-number", default="",
                        help="Case number for the report")
    parser.add_argument("--gdrive-folder", default=None,
                        help="Google Drive folder URL or ID to scan for documents")
    parser.add_argument("--gdrive-links-col", default=None,
                        help="Column name in data containing Google Drive links")
    parser.add_argument("--compare", nargs=2, default=None, metavar=("DOC_A", "DOC_B"),
                        help="Two documents/URLs to compare side-by-side")
    parser.add_argument("--output", "-o", default="output",
                        help="Output directory (default: output)")
    parser.add_argument("--interactive", action="store_true",
                        help="Run in interactive mode")

    args = parser.parse_args()

    if args.interactive:
        print_header()
        args.input = input("Input path (file or directory) [./input]: ").strip() or "./input"
        args.state = input("State code (e.g., OH) []: ").strip() or None
        args.county = input("County name []: ").strip() or None
        args.case_name = input("Case name []: ").strip() or ""
        args.case_number = input("Case number []: ").strip() or ""
        gdrive = input("Google Drive folder URL (or press Enter to skip) []: ").strip()
        args.gdrive_folder = gdrive if gdrive else None
        gdrive_col = input("Column with Drive links (or press Enter to skip) []: ").strip()
        args.gdrive_links_col = gdrive_col if gdrive_col else None
        compare = input("Compare two docs? Enter two paths separated by space (or Enter to skip): ").strip()
        args.compare = compare.split()[:2] if compare else None

    run_analysis(
        input_path=args.input,
        state_code=args.state,
        county=args.county,
        gdrive_folder=args.gdrive_folder,
        gdrive_links_col=args.gdrive_links_col,
        case_name=args.case_name,
        case_number=args.case_number,
        compare_docs=args.compare,
        output_dir=args.output,
    )


if __name__ == "__main__":
    main()
