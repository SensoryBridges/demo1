"""
Google Drive Integration Module.

Connects to Google Drive to search, read, and extract data from
court-related documents and evidence files.

Setup:
1. Go to https://console.cloud.google.com/
2. Create a project and enable Google Drive API + Google Docs API + Google Sheets API
3. Create OAuth 2.0 credentials (Desktop application)
4. Download credentials.json to the project root directory
5. On first run, a browser window will open for authorization
"""

import os
import io
import re
import json
import mimetypes
from pathlib import Path

# Conditional imports — graceful fallback if not installed
try:
    from google.oauth2.credentials import Credentials
    from google_auth_oauthlib.flow import InstalledAppFlow
    from google.auth.transport.requests import Request
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaIoBaseDownload
    HAS_GOOGLE = True
except ImportError:
    HAS_GOOGLE = False

try:
    import fitz as pymupdf  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# Scopes required for read-only Drive access + Docs/Sheets read
SCOPES = [
    "https://www.googleapis.com/auth/drive.readonly",
    "https://www.googleapis.com/auth/documents.readonly",
    "https://www.googleapis.com/auth/spreadsheets.readonly",
]

# Court document keyword patterns for smart searching
COURT_DOC_KEYWORDS = [
    "motion", "order", "petition", "response", "filing", "affidavit",
    "declaration", "exhibit", "judgment", "decree", "stipulation",
    "memorandum", "brief", "subpoena", "summons", "complaint",
    "custody", "visitation", "parenting", "support", "alimony",
    "contempt", "modification", "appeal", "guardian", "GAL",
    "evaluation", "report", "hearing", "transcript", "deposition",
    "discovery", "interrogator", "docket", "notice", "certificate",
]


class GoogleDriveConnector:
    """Handles authentication and file operations with Google Drive."""

    def __init__(self, credentials_path="credentials.json", token_path="token.json"):
        self.credentials_path = credentials_path
        self.token_path = token_path
        self.creds = None
        self.drive_service = None
        self.docs_service = None
        self.sheets_service = None
        self._authenticated = False

    def authenticate(self):
        """Authenticate with Google APIs using OAuth 2.0."""
        if not HAS_GOOGLE:
            raise ImportError(
                "Google API libraries not installed. Run:\n"
                "  pip install google-auth google-auth-oauthlib "
                "google-auth-httplib2 google-api-python-client"
            )

        creds = None

        # Load existing token
        if os.path.exists(self.token_path):
            creds = Credentials.from_authorized_user_file(self.token_path, SCOPES)

        # Refresh or create new credentials
        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token:
                creds.refresh(Request())
            else:
                if not os.path.exists(self.credentials_path):
                    raise FileNotFoundError(
                        f"Credentials file not found: {self.credentials_path}\n"
                        "Download OAuth 2.0 credentials from Google Cloud Console:\n"
                        "  https://console.cloud.google.com/apis/credentials\n"
                        "Save as 'credentials.json' in the project root."
                    )
                flow = InstalledAppFlow.from_client_secrets_file(
                    self.credentials_path, SCOPES
                )
                creds = flow.run_local_server(port=0)

            # Save token for future runs
            with open(self.token_path, "w") as token_file:
                token_file.write(creds.to_json())

        self.creds = creds
        self.drive_service = build("drive", "v3", credentials=creds)
        self.docs_service = build("docs", "v1", credentials=creds)
        self.sheets_service = build("sheets", "v4", credentials=creds)
        self._authenticated = True
        return True

    def _require_auth(self):
        if not self._authenticated:
            raise RuntimeError("Not authenticated. Call authenticate() first.")

    # ------------------------------------------------------------------
    # Search
    # ------------------------------------------------------------------

    def search_files(self, query=None, folder_id=None, mime_types=None,
                     max_results=100):
        """
        Search for files in Google Drive.

        Args:
            query: Free-text search string
            folder_id: Restrict to a specific folder
            mime_types: List of MIME types to filter
            max_results: Maximum number of results

        Returns:
            List of file metadata dicts
        """
        self._require_auth()

        q_parts = ["trashed = false"]
        if folder_id:
            q_parts.append(f"'{folder_id}' in parents")
        if query:
            q_parts.append(f"fullText contains '{query}'")
        if mime_types:
            mime_q = " or ".join(f"mimeType='{m}'" for m in mime_types)
            q_parts.append(f"({mime_q})")

        q_string = " and ".join(q_parts)

        results = []
        page_token = None
        while len(results) < max_results:
            resp = self.drive_service.files().list(
                q=q_string,
                pageSize=min(100, max_results - len(results)),
                pageToken=page_token,
                fields="nextPageToken, files(id, name, mimeType, modifiedTime, "
                       "size, parents, webViewLink, description)",
                orderBy="modifiedTime desc",
            ).execute()

            results.extend(resp.get("files", []))
            page_token = resp.get("nextPageToken")
            if not page_token:
                break

        return results

    def search_court_documents(self, folder_id=None, max_results=200):
        """Search specifically for court-related documents."""
        all_files = []
        seen_ids = set()

        # Search using court keywords
        for keyword in COURT_DOC_KEYWORDS[:15]:  # Top keywords
            files = self.search_files(
                query=keyword,
                folder_id=folder_id,
                max_results=20,
            )
            for f in files:
                if f["id"] not in seen_ids:
                    seen_ids.add(f["id"])
                    f["matched_keyword"] = keyword
                    all_files.append(f)

            if len(all_files) >= max_results:
                break

        return all_files[:max_results]

    def list_folder(self, folder_id, recursive=False):
        """List all files in a folder, optionally recursive."""
        self._require_auth()
        files = []
        q = f"'{folder_id}' in parents and trashed = false"
        page_token = None

        while True:
            resp = self.drive_service.files().list(
                q=q, pageSize=100, pageToken=page_token,
                fields="nextPageToken, files(id, name, mimeType, modifiedTime, "
                       "size, parents, webViewLink)",
            ).execute()
            batch = resp.get("files", [])
            files.extend(batch)

            if recursive:
                for f in batch:
                    if f["mimeType"] == "application/vnd.google-apps.folder":
                        files.extend(self.list_folder(f["id"], recursive=True))

            page_token = resp.get("nextPageToken")
            if not page_token:
                break

        return files

    def extract_folder_id(self, url_or_id):
        """Extract a Google Drive folder ID from a URL or return as-is."""
        if "/" in url_or_id:
            match = re.search(r"/folders/([a-zA-Z0-9_-]+)", url_or_id)
            if match:
                return match.group(1)
            match = re.search(r"id=([a-zA-Z0-9_-]+)", url_or_id)
            if match:
                return match.group(1)
        return url_or_id

    def extract_file_id(self, url_or_id):
        """Extract a Google Drive file ID from a URL or return as-is."""
        if "/" in url_or_id:
            match = re.search(r"/d/([a-zA-Z0-9_-]+)", url_or_id)
            if match:
                return match.group(1)
            match = re.search(r"id=([a-zA-Z0-9_-]+)", url_or_id)
            if match:
                return match.group(1)
        return url_or_id

    # ------------------------------------------------------------------
    # Read / Download
    # ------------------------------------------------------------------

    def read_google_doc(self, file_id):
        """Read a Google Docs document and return structured text."""
        self._require_auth()
        file_id = self.extract_file_id(file_id)
        doc = self.docs_service.documents().get(documentId=file_id).execute()
        title = doc.get("title", "Untitled")
        body = doc.get("body", {})
        text_parts = []

        for element in body.get("content", []):
            if "paragraph" in element:
                for pe in element["paragraph"].get("elements", []):
                    if "textRun" in pe:
                        text_parts.append(pe["textRun"]["content"])
            elif "table" in element:
                for row in element["table"].get("tableRows", []):
                    row_text = []
                    for cell in row.get("tableCells", []):
                        cell_text = ""
                        for p in cell.get("content", []):
                            if "paragraph" in p:
                                for pe in p["paragraph"].get("elements", []):
                                    if "textRun" in pe:
                                        cell_text += pe["textRun"]["content"]
                        row_text.append(cell_text.strip())
                    text_parts.append("\t".join(row_text))

        return {"title": title, "text": "".join(text_parts), "file_id": file_id}

    def read_google_sheet(self, file_id, range_name=None):
        """Read a Google Sheets spreadsheet and return as DataFrame(s)."""
        self._require_auth()
        import pandas as pd
        file_id = self.extract_file_id(file_id)

        # Get sheet metadata
        meta = self.sheets_service.spreadsheets().get(
            spreadsheetId=file_id
        ).execute()
        sheet_names = [s["properties"]["title"] for s in meta.get("sheets", [])]

        dataframes = {}
        for sheet in sheet_names:
            r = range_name if range_name else sheet
            result = self.sheets_service.spreadsheets().values().get(
                spreadsheetId=file_id, range=r
            ).execute()
            values = result.get("values", [])
            if values:
                headers = [str(h).strip().lower().replace(" ", "_") for h in values[0]]
                data = values[1:]
                # Pad rows to header length
                data = [row + [""] * (len(headers) - len(row)) for row in data]
                df = pd.DataFrame(data, columns=headers)
                dataframes[sheet] = df

        return dataframes

    def download_file(self, file_id, output_dir="downloads"):
        """Download a file from Drive to local disk."""
        self._require_auth()
        file_id = self.extract_file_id(file_id)
        os.makedirs(output_dir, exist_ok=True)

        # Get file metadata
        meta = self.drive_service.files().get(
            fileId=file_id, fields="name, mimeType"
        ).execute()
        name = meta["name"]
        mime = meta["mimeType"]

        # Google Workspace files need export
        export_map = {
            "application/vnd.google-apps.document": (
                "application/pdf", ".pdf"
            ),
            "application/vnd.google-apps.spreadsheet": (
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                ".xlsx"
            ),
            "application/vnd.google-apps.presentation": (
                "application/pdf", ".pdf"
            ),
        }

        if mime in export_map:
            export_mime, ext = export_map[mime]
            request = self.drive_service.files().export_media(
                fileId=file_id, mimeType=export_mime
            )
            if not name.endswith(ext):
                name += ext
        else:
            request = self.drive_service.files().get_media(fileId=file_id)

        filepath = os.path.join(output_dir, name)
        with io.FileIO(filepath, "wb") as fh:
            downloader = MediaIoBaseDownload(fh, request)
            done = False
            while not done:
                _, done = downloader.next_chunk()

        return filepath

    def read_file_text(self, file_id):
        """Download and extract text from any supported file type."""
        self._require_auth()
        file_id = self.extract_file_id(file_id)

        meta = self.drive_service.files().get(
            fileId=file_id, fields="name, mimeType"
        ).execute()
        mime = meta["mimeType"]
        name = meta["name"]

        # Google Docs — use API directly
        if mime == "application/vnd.google-apps.document":
            return self.read_google_doc(file_id)

        # Google Sheets — use API directly
        if mime == "application/vnd.google-apps.spreadsheet":
            dfs = self.read_google_sheet(file_id)
            text = ""
            for sheet_name, df in dfs.items():
                text += f"\n=== {sheet_name} ===\n"
                text += df.to_string(index=False) + "\n"
            return {"title": name, "text": text, "file_id": file_id}

        # Binary files — download and extract
        filepath = self.download_file(file_id, output_dir="/tmp/drive_downloads")
        text = self._extract_text_from_file(filepath)
        return {"title": name, "text": text, "file_id": file_id, "local_path": filepath}

    def _extract_text_from_file(self, filepath):
        """Extract text from a local file based on its type."""
        ext = os.path.splitext(filepath)[1].lower()

        if ext == ".pdf":
            return self._extract_pdf_text(filepath)
        elif ext in (".docx", ".doc"):
            return self._extract_docx_text(filepath)
        elif ext in (".txt", ".csv", ".tsv", ".log"):
            with open(filepath, "r", errors="ignore") as f:
                return f.read()
        elif ext in (".xlsx", ".xls"):
            return self._extract_excel_text(filepath)
        elif ext in (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff"):
            return f"[Image file: {os.path.basename(filepath)} — OCR not available. " \
                   f"Install pytesseract for image text extraction.]"
        else:
            return f"[Unsupported file type: {ext}]"

    def _extract_pdf_text(self, filepath):
        """Extract text from a PDF file."""
        if HAS_PYMUPDF:
            doc = pymupdf.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text() + "\n"
            doc.close()
            return text
        else:
            return ("[PDF text extraction requires PyMuPDF. "
                    "Install with: pip install pymupdf]")

    def _extract_docx_text(self, filepath):
        """Extract text from a Word document."""
        if HAS_DOCX:
            doc = DocxDocument(filepath)
            paragraphs = [p.text for p in doc.paragraphs]
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    paragraphs.append("\t".join(row_text))
            return "\n".join(paragraphs)
        else:
            return "[DOCX extraction requires python-docx. Install with: pip install python-docx]"

    def _extract_excel_text(self, filepath):
        """Extract text from an Excel file."""
        import pandas as pd
        sheets = pd.read_excel(filepath, sheet_name=None, engine="openpyxl")
        text = ""
        for name, df in sheets.items():
            text += f"\n=== {name} ===\n"
            text += df.to_string(index=False) + "\n"
        return text

    # ------------------------------------------------------------------
    # Batch operations for court case analysis
    # ------------------------------------------------------------------

    def gather_case_documents(self, folder_id_or_url, output_dir="case_documents"):
        """
        Gather all documents from a case folder, extract text,
        and return structured results.
        """
        folder_id = self.extract_folder_id(folder_id_or_url)
        files = self.list_folder(folder_id, recursive=True)
        os.makedirs(output_dir, exist_ok=True)

        documents = []
        for f in files:
            if f["mimeType"] == "application/vnd.google-apps.folder":
                continue

            print(f"  Processing: {f['name']}")
            try:
                result = self.read_file_text(f["id"])
                doc_entry = {
                    "file_id": f["id"],
                    "name": f["name"],
                    "mime_type": f["mimeType"],
                    "modified": f.get("modifiedTime", ""),
                    "link": f.get("webViewLink", ""),
                    "text": result.get("text", ""),
                    "text_length": len(result.get("text", "")),
                }
                documents.append(doc_entry)
            except Exception as e:
                documents.append({
                    "file_id": f["id"],
                    "name": f["name"],
                    "error": str(e),
                    "text": "",
                })

        return documents

    def resolve_drive_links(self, df, link_column="link"):
        """
        Given a DataFrame with a column of Google Drive links,
        resolve each link and extract the document text.
        """
        if link_column not in df.columns:
            return df

        texts = []
        titles = []
        for _, row in df.iterrows():
            link = row[link_column]
            if not isinstance(link, str) or "drive.google" not in link:
                texts.append("")
                titles.append("")
                continue

            try:
                file_id = self.extract_file_id(link)
                result = self.read_file_text(file_id)
                texts.append(result.get("text", ""))
                titles.append(result.get("title", ""))
            except Exception as e:
                texts.append(f"[Error: {e}]")
                titles.append("")

        df["document_text"] = texts
        df["document_title"] = titles
        return df
