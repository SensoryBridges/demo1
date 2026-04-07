"""
Report Generation Module.

Produces legally defensible reports in Excel and Word formats
with full analysis, charts, evidence citations, and methodology documentation.
"""

import os
from datetime import datetime

import pandas as pd
import numpy as np
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side, numbers
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.chart import BarChart, PieChart, Reference
from openpyxl.formatting.rule import CellIsRule, ColorScaleRule, DataBarRule
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


class ExcelReportGenerator:
    """Generates comprehensive Excel workbook reports."""

    HEADER_FILL = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    HEADER_FONT = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    ALERT_FILL = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
    GOOD_FILL = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
    WARN_FILL = PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid")
    CRITICAL_FILL = PatternFill(start_color="C00000", end_color="C00000", fill_type="solid")
    CRITICAL_FONT = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    HIGH_FILL = PatternFill(start_color="FF6B6B", end_color="FF6B6B", fill_type="solid")
    MODERATE_FILL = PatternFill(start_color="FFD93D", end_color="FFD93D", fill_type="solid")
    LOW_FILL = PatternFill(start_color="A8E6CF", end_color="A8E6CF", fill_type="solid")
    SUBHEADER_FILL = PatternFill(start_color="D6E4F0", end_color="D6E4F0", fill_type="solid")
    SUBHEADER_FONT = Font(name="Calibri", size=11, bold=True, color="1F4E79")
    BORDER = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin")
    )

    def __init__(self, output_path="output/analysis_report.xlsx"):
        self.output_path = output_path
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        self.wb = openpyxl.Workbook()
        self.wb.remove(self.wb.active)

    def _add_sheet(self, name, df=None, title=None):
        """Add a formatted sheet with optional DataFrame data."""
        ws = self.wb.create_sheet(title=name[:31])

        row_offset = 1
        if title:
            ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max(len(df.columns) if df is not None else 5, 5))
            cell = ws.cell(row=1, column=1, value=title)
            cell.font = Font(name="Calibri", size=14, bold=True, color="1F4E79")
            cell.alignment = Alignment(horizontal="center")
            row_offset = 3

        if df is not None and not df.empty:
            # Headers
            for col_idx, col_name in enumerate(df.columns, 1):
                cell = ws.cell(row=row_offset, column=col_idx, value=str(col_name).replace("_", " ").title())
                cell.fill = self.HEADER_FILL
                cell.font = self.HEADER_FONT
                cell.alignment = Alignment(horizontal="center", wrap_text=True)
                cell.border = self.BORDER

            # Data
            for row_idx, row_data in enumerate(df.itertuples(index=False), row_offset + 1):
                for col_idx, value in enumerate(row_data, 1):
                    cell = ws.cell(row=row_idx, column=col_idx)
                    if isinstance(value, (np.integer,)):
                        cell.value = int(value)
                    elif isinstance(value, (np.floating,)):
                        cell.value = round(float(value), 3)
                    elif pd.isna(value):
                        cell.value = ""
                    else:
                        cell.value = str(value) if not isinstance(value, (int, float, datetime)) else value
                    cell.border = self.BORDER
                    cell.alignment = Alignment(wrap_text=True, vertical="top")

            # Auto-width columns
            for col_idx in range(1, len(df.columns) + 1):
                max_len = max(
                    len(str(ws.cell(row=r, column=col_idx).value or ""))
                    for r in range(row_offset, min(row_offset + len(df) + 1, row_offset + 50))
                )
                ws.column_dimensions[openpyxl.utils.get_column_letter(col_idx)].width = min(50, max(12, max_len + 2))

        return ws

    def _add_dict_sheet(self, name, data_dict, title=None):
        """Add a sheet from a dictionary."""
        rows = []
        self._flatten_dict(data_dict, rows)
        df = pd.DataFrame(rows, columns=["Field", "Value"])
        return self._add_sheet(name, df, title)

    def _flatten_dict(self, d, rows, prefix=""):
        """Flatten nested dict into rows."""
        for key, value in d.items():
            label = f"{prefix}{key}".replace("_", " ").title()
            if isinstance(value, dict):
                rows.append([label, "---"])
                self._flatten_dict(value, rows, prefix=f"  ")
            elif isinstance(value, list):
                rows.append([label, f"({len(value)} items)"])
                for i, item in enumerate(value[:50]):
                    if isinstance(item, dict):
                        for k, v in item.items():
                            rows.append([f"    [{i + 1}] {k}", str(v)[:500]])
                    else:
                        rows.append([f"    [{i + 1}]", str(item)[:500]])
            else:
                rows.append([label, str(value)[:500]])

    def add_executive_summary(self, summary_data):
        """Add executive summary sheet."""
        self._add_dict_sheet("Executive Summary", summary_data,
                             "CASE ANALYSIS — EXECUTIVE SUMMARY")

    def add_abuse_analysis(self, classified_df):
        """Add abuse classification results."""
        self._add_sheet("Abuse Analysis", classified_df,
                        "POST-SEPARATION ABUSE ANALYSIS")

    def add_cycle_predictions(self, predictions_data):
        """Add cycle prediction results."""
        if "predictions" in predictions_data:
            pred_df = pd.DataFrame(predictions_data["predictions"])
            self._add_sheet("Cycle Predictions", pred_df,
                            "ABUSE CYCLE PREDICTIONS")
        if "analysis_summary" in predictions_data:
            self._add_dict_sheet("Pattern Analysis", predictions_data["analysis_summary"],
                                 "PATTERN ANALYSIS DETAILS")

    def add_court_timeline(self, timeline_data):
        """Add court case timeline."""
        if isinstance(timeline_data, list):
            df = pd.DataFrame(timeline_data)
        else:
            df = timeline_data
        self._add_sheet("Court Timeline", df, "COURT CASE TIMELINE")

    def add_officials_analysis(self, officials_df, earnings_data=None):
        """Add court officials analysis."""
        self._add_sheet("Officials", officials_df, "COURT OFFICIALS ANALYSIS")
        if earnings_data:
            earn_df = pd.DataFrame.from_dict(earnings_data, orient="index")
            self._add_sheet("Official Earnings", earn_df.reset_index().rename(columns={"index": "name"}),
                            "ESTIMATED OFFICIAL EARNINGS")

    def add_violations(self, violations):
        """Add procedural violations sheet."""
        if violations:
            df = pd.DataFrame(violations)
            ws = self._add_sheet("Violations", df, "PROCEDURAL VIOLATIONS IDENTIFIED")
            # Highlight severity
            if "severity" in df.columns:
                sev_col = list(df.columns).index("severity") + 1
                for row_idx in range(4, 4 + len(df)):
                    cell = ws.cell(row=row_idx, column=sev_col)
                    if "HIGH" in str(cell.value) or "CRITICAL" in str(cell.value):
                        cell.fill = self.ALERT_FILL
                    elif "MODERATE" in str(cell.value):
                        cell.fill = self.WARN_FILL

    def add_financial_analysis(self, costs_data):
        """Add financial analysis sheet."""
        self._add_dict_sheet("Financial Analysis", costs_data,
                             "FINANCIAL ANALYSIS")

    def add_connections(self, connections):
        """Add connections analysis."""
        if connections:
            df = pd.DataFrame(connections)
            self._add_sheet("Connections", df, "OFFICIAL CONNECTIONS ANALYSIS")

    def add_scorecards(self, scorecards_df):
        """Add official scorecards."""
        ws = self._add_sheet("Scorecards", scorecards_df, "COURT OFFICIAL SCORECARDS")
        # Highlight grades
        if not scorecards_df.empty and "overall_score" in scorecards_df.columns:
            score_col = list(scorecards_df.columns).index("overall_score") + 1
            for row_idx in range(4, 4 + len(scorecards_df)):
                cell = ws.cell(row=row_idx, column=score_col)
                try:
                    val = float(cell.value)
                    if val < 50:
                        cell.fill = self.ALERT_FILL
                    elif val < 70:
                        cell.fill = self.WARN_FILL
                    else:
                        cell.fill = self.GOOD_FILL
                except (ValueError, TypeError):
                    pass

    def add_document_comparison(self, comparison_result):
        """Add document comparison side-by-side view."""
        sbs = comparison_result.get("side_by_side", [])
        if sbs:
            df = pd.DataFrame(sbs)
            ws = self._add_sheet("Doc Comparison", df,
                                 "DOCUMENT COMPARISON — SIDE BY SIDE")
            # Highlight differences
            if "flag" in df.columns:
                flag_col = list(df.columns).index("flag") + 1
                for row_idx in range(4, 4 + len(df)):
                    cell = ws.cell(row=row_idx, column=flag_col)
                    if cell.value == "DIFFERENT":
                        cell.fill = self.ALERT_FILL
                    elif cell.value == "SIMILAR":
                        cell.fill = self.WARN_FILL
                    elif cell.value == "MATCH":
                        cell.fill = self.GOOD_FILL

        # Contradictions
        contras = comparison_result.get("contradictions", [])
        if contras:
            df2 = pd.DataFrame(contras)
            self._add_sheet("Contradictions", df2, "DOCUMENT CONTRADICTIONS")

    def add_formula_documentation(self, formula_doc):
        """Add CASI formula documentation."""
        self._add_dict_sheet("CASI Formula", formula_doc,
                             "COMPOSITE ABUSE SEVERITY INDEX (CASI) — FORMULA DOCUMENTATION")

    def add_case_outcome_prediction(self, prediction_data):
        """Add outcome prediction results."""
        self._add_dict_sheet("Outcome Prediction", prediction_data,
                             "CASE OUTCOME PREDICTION")

    def add_methodology(self):
        """Add methodology documentation sheet for legal defensibility."""
        methodology = {
            "Report Purpose": "Forensic analysis of post-separation abuse patterns and court case proceedings",
            "Methodology": {
                "Data Collection": "All data sourced from court records, documented communications, and official filings",
                "Classification System": "Incidents classified using standardized abuse taxonomy aligned with VAWA definitions",
                "Statistical Methods": "Time-series analysis, correlation testing, linear regression, peak detection",
                "Prediction Model": "Composite Abuse Severity Index (CASI) — see formula documentation sheet",
                "Validation": "All statistical findings include p-values and confidence intervals",
            },
            "Legal Standards": {
                "Evidence Standard": "Analysis based on documented evidence only — no conjecture",
                "Statistical Significance": "Findings marked significant at p < 0.05 level",
                "Reproducibility": "All calculations are deterministic and reproducible from source data",
                "Limitations": "Predictions are probabilistic projections, not certainties. "
                              "Analysis quality depends on completeness of input data.",
            },
            "Disclaimers": {
                "Not Legal Advice": "This report is an analytical tool and does not constitute legal advice",
                "Professional Review": "Findings should be reviewed by qualified legal and mental health professionals",
                "Data Dependency": "Conclusions are only as reliable as the underlying data",
            },
            "Report Generated": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "Tool Version": "Abuse Court Analyzer v1.0.0",
        }
        self._add_dict_sheet("Methodology", methodology,
                             "METHODOLOGY & LEGAL DEFENSIBILITY")

    def _apply_autofilter(self, ws, header_row, num_cols):
        """Apply auto-filter to all columns starting at the given header row."""
        last_col_letter = openpyxl.utils.get_column_letter(num_cols)
        ws.auto_filter.ref = f"A{header_row}:{last_col_letter}{ws.max_row}"

    def _apply_severity_coloring(self, ws, df, severity_col_name, data_start_row):
        """Apply color coding to cells in a severity column."""
        if severity_col_name not in df.columns:
            return
        col_idx = list(df.columns).index(severity_col_name) + 1
        for row_idx in range(data_start_row, data_start_row + len(df)):
            cell = ws.cell(row=row_idx, column=col_idx)
            val = str(cell.value).upper()
            if "CRITICAL" in val:
                cell.fill = self.CRITICAL_FILL
                cell.font = self.CRITICAL_FONT
            elif "HIGH" in val or "SEVERE" in val:
                cell.fill = self.HIGH_FILL
            elif "MODERATE" in val or "MEDIUM" in val:
                cell.fill = self.MODERATE_FILL
            elif "LOW" in val or "MINOR" in val:
                cell.fill = self.LOW_FILL
            elif "GOOD" in val or "COMPLIANT" in val:
                cell.fill = self.GOOD_FILL

    def _apply_row_severity_coloring(self, ws, df, severity_col_name, data_start_row, num_cols):
        """Apply color coding to entire rows based on a severity column value."""
        if severity_col_name not in df.columns:
            return
        col_idx = list(df.columns).index(severity_col_name) + 1
        for row_idx in range(data_start_row, data_start_row + len(df)):
            sev_cell = ws.cell(row=row_idx, column=col_idx)
            val = str(sev_cell.value).upper()
            fill = None
            font = None
            if "CRITICAL" in val:
                fill = self.CRITICAL_FILL
                font = self.CRITICAL_FONT
            elif "HIGH" in val or "SEVERE" in val:
                fill = self.HIGH_FILL
            elif "MODERATE" in val or "MEDIUM" in val:
                fill = self.MODERATE_FILL
            if fill:
                for c in range(1, num_cols + 1):
                    ws.cell(row=row_idx, column=c).fill = fill
                    if font:
                        ws.cell(row=row_idx, column=c).font = font

    def add_detailed_event_timeline(self, timeline):
        """Add detailed event timeline sheet with every event color-coded by severity.

        Args:
            timeline: list of dicts with keys: date, description, persons_involved,
                      laws_broken, human_rights_violations, method, source_document,
                      severity.
        """
        if not timeline:
            return

        columns = [
            "date", "description", "persons_involved", "laws_broken",
            "human_rights_violations", "method", "source_document", "severity"
        ]
        rows = []
        for event in timeline:
            row = {}
            for col in columns:
                val = event.get(col, "")
                if isinstance(val, list):
                    val = "; ".join(str(v) for v in val)
                row[col] = val
            rows.append(row)

        df = pd.DataFrame(rows, columns=columns)

        # Sort by date when possible
        try:
            df["date"] = pd.to_datetime(df["date"], errors="coerce")
            df = df.sort_values("date", na_position="last")
        except Exception:
            pass

        ws = self._add_sheet("Event Timeline", df, "DETAILED EVENT TIMELINE")

        # title=row 1, blank=row 2, header=row 3, data starts at row 4
        data_start_row = 4
        num_cols = len(columns)

        self._apply_autofilter(ws, 3, num_cols)
        self._apply_row_severity_coloring(ws, df, "severity", data_start_row, num_cols)

        # Format date column
        for row_idx in range(data_start_row, data_start_row + len(df)):
            cell = ws.cell(row=row_idx, column=1)
            if isinstance(cell.value, datetime):
                cell.number_format = "YYYY-MM-DD"

        # Set appropriate column widths for content-heavy columns
        ws.column_dimensions["B"].width = 60
        ws.column_dimensions["C"].width = 30
        ws.column_dimensions["D"].width = 40
        ws.column_dimensions["E"].width = 40
        ws.column_dimensions["F"].width = 25
        ws.column_dimensions["G"].width = 30

        # Freeze header row so filters remain visible while scrolling
        ws.freeze_panes = "A4"

    def add_365_day_analysis(self, analysis):
        """Add 365-day order compliance analysis sheet with timeline.

        Args:
            analysis: dict with keys: orders (list of dicts with order_date,
                      order_type, compliance_deadline, compliance_status,
                      days_elapsed, violation_details, legal_basis),
                      summary (dict), timeline (list).
        """
        if not analysis:
            return

        orders = analysis.get("orders", [])
        if orders:
            order_cols = [
                "order_date", "order_type", "compliance_deadline",
                "compliance_status", "days_elapsed", "violation_details", "legal_basis"
            ]
            rows = []
            for order in orders:
                row = {col: order.get(col, "") for col in order_cols}
                if isinstance(row.get("violation_details"), list):
                    row["violation_details"] = "; ".join(str(v) for v in row["violation_details"])
                rows.append(row)

            df = pd.DataFrame(rows, columns=order_cols)
            ws = self._add_sheet("365-Day Analysis", df,
                                 "365-DAY ORDER COMPLIANCE ANALYSIS")
            data_start_row = 4
            num_cols = len(order_cols)

            self._apply_autofilter(ws, 3, num_cols)
            self._apply_severity_coloring(ws, df, "compliance_status", data_start_row)

            # Format date columns
            for col_name in ["order_date", "compliance_deadline"]:
                if col_name in df.columns:
                    col_idx = list(df.columns).index(col_name) + 1
                    for row_idx in range(data_start_row, data_start_row + len(df)):
                        ws.cell(row=row_idx, column=col_idx).number_format = "YYYY-MM-DD"

            # Format days_elapsed as integer
            if "days_elapsed" in df.columns:
                col_idx = list(df.columns).index("days_elapsed") + 1
                for row_idx in range(data_start_row, data_start_row + len(df)):
                    ws.cell(row=row_idx, column=col_idx).number_format = "0"

            ws.column_dimensions["B"].width = 30
            ws.column_dimensions["F"].width = 50
            ws.column_dimensions["G"].width = 40
            ws.freeze_panes = "A4"

        # Summary sub-sheet
        summary = analysis.get("summary", {})
        if summary:
            self._add_dict_sheet("365-Day Summary", summary,
                                 "365-DAY COMPLIANCE SUMMARY")

    def add_uccjea_analysis(self, analysis):
        """Add UCCJEA jurisdiction analysis sheet.

        Args:
            analysis: dict with keys: jurisdiction_findings (list of dicts),
                      address_history (list of dicts), home_state_determination (dict),
                      significant_connection (dict), dismissal_grounds (dict),
                      inconvenient_forum_factors (dict).
        """
        if not analysis:
            return

        # Jurisdiction findings
        findings = analysis.get("jurisdiction_findings", [])
        if findings:
            df = pd.DataFrame(findings)
            ws = self._add_sheet("UCCJEA Analysis", df,
                                 "UCCJEA JURISDICTION ANALYSIS")
            data_start_row = 4
            self._apply_autofilter(ws, 3, len(df.columns))
            self._apply_severity_coloring(ws, df, "status", data_start_row)
            ws.freeze_panes = "A4"

        # Address history
        addr_history = analysis.get("address_history", [])
        if addr_history:
            addr_df = pd.DataFrame(addr_history)
            ws2 = self._add_sheet("Address History", addr_df,
                                  "CHILD ADDRESS HISTORY — UCCJEA")
            self._apply_autofilter(ws2, 3, len(addr_df.columns))

            for col_name in addr_df.columns:
                if "date" in col_name.lower():
                    col_idx = list(addr_df.columns).index(col_name) + 1
                    for row_idx in range(4, 4 + len(addr_df)):
                        ws2.cell(row=row_idx, column=col_idx).number_format = "YYYY-MM-DD"
            ws2.freeze_panes = "A4"

        # Dismissal grounds summary
        dismissal = analysis.get("dismissal_grounds", {})
        if dismissal:
            self._add_dict_sheet("UCCJEA Dismissal", dismissal,
                                 "UCCJEA DISMISSAL GROUNDS")

    def add_best_interest_factors(self, analysis):
        """Add best-interest-of-child 17-factor analysis sheet.

        Args:
            analysis: dict with keys: factors (list of 17 dicts, each with
                      factor_name, factor_number, parent_a_score, parent_b_score,
                      parent_a_evidence, parent_b_evidence, weight),
                      parent_a_name, parent_b_name, summary (dict).
        """
        if not analysis:
            return

        factors = analysis.get("factors", [])
        parent_a = analysis.get("parent_a_name", "Parent A")
        parent_b = analysis.get("parent_b_name", "Parent B")

        if not factors:
            return

        columns = [
            "factor_number", "factor_name",
            f"{parent_a}_score", f"{parent_b}_score",
            f"{parent_a}_evidence", f"{parent_b}_evidence",
            "weight"
        ]
        rows = []
        for f in factors:
            rows.append({
                "factor_number": f.get("factor_number", ""),
                "factor_name": f.get("factor_name", ""),
                f"{parent_a}_score": f.get("parent_a_score", ""),
                f"{parent_b}_score": f.get("parent_b_score", ""),
                f"{parent_a}_evidence": f.get("parent_a_evidence", ""),
                f"{parent_b}_evidence": f.get("parent_b_evidence", ""),
                "weight": f.get("weight", 1.0),
            })

        df = pd.DataFrame(rows, columns=columns)
        ws = self._add_sheet("Best Interest Factors", df,
                             "BEST INTEREST OF THE CHILD — 17-FACTOR ANALYSIS")

        data_start_row = 4
        num_cols = len(columns)
        self._apply_autofilter(ws, 3, num_cols)

        # Conditional formatting on score columns (1-10 scale: green=high, red=low)
        score_a_col_idx = 3
        score_b_col_idx = 4
        score_a_letter = openpyxl.utils.get_column_letter(score_a_col_idx)
        score_b_letter = openpyxl.utils.get_column_letter(score_b_col_idx)
        last_data_row = data_start_row + len(df) - 1

        for col_letter in [score_a_letter, score_b_letter]:
            cell_range = f"{col_letter}{data_start_row}:{col_letter}{last_data_row}"
            ws.conditional_formatting.add(
                cell_range,
                CellIsRule(operator="lessThan", formula=["4"],
                           fill=self.ALERT_FILL)
            )
            ws.conditional_formatting.add(
                cell_range,
                CellIsRule(operator="between", formula=["4", "6"],
                           fill=self.WARN_FILL)
            )
            ws.conditional_formatting.add(
                cell_range,
                CellIsRule(operator="greaterThan", formula=["6"],
                           fill=self.GOOD_FILL)
            )

        # Format weight column as percentage when values are <= 1
        weight_col_idx = list(df.columns).index("weight") + 1
        for row_idx in range(data_start_row, data_start_row + len(df)):
            cell = ws.cell(row=row_idx, column=weight_col_idx)
            try:
                val = float(cell.value)
                if val <= 1.0:
                    cell.number_format = "0.0%"
            except (ValueError, TypeError):
                pass

        # Widen evidence columns
        ws.column_dimensions["B"].width = 40
        ws.column_dimensions["E"].width = 55
        ws.column_dimensions["F"].width = 55

        # Add weighted-total row
        total_row = data_start_row + len(df)
        ws.cell(row=total_row, column=1, value="").border = self.BORDER
        total_cell = ws.cell(row=total_row, column=2, value="WEIGHTED TOTAL")
        total_cell.font = Font(name="Calibri", size=11, bold=True)
        total_cell.border = self.BORDER

        for col_idx in [score_a_col_idx, score_b_col_idx]:
            col_letter = openpyxl.utils.get_column_letter(col_idx)
            weight_letter = openpyxl.utils.get_column_letter(weight_col_idx)
            formula = (
                f"=SUMPRODUCT({col_letter}{data_start_row}:{col_letter}{last_data_row},"
                f"{weight_letter}{data_start_row}:{weight_letter}{last_data_row})"
                f"/SUM({weight_letter}{data_start_row}:{weight_letter}{last_data_row})"
            )
            cell = ws.cell(row=total_row, column=col_idx, value=formula)
            cell.font = Font(name="Calibri", size=11, bold=True)
            cell.number_format = "0.00"
            cell.border = self.BORDER
            cell.fill = self.SUBHEADER_FILL

        ws.freeze_panes = "A4"

        # Summary sub-sheet
        summary = analysis.get("summary", {})
        if summary:
            self._add_dict_sheet("Best Interest Summary", summary,
                                 "BEST INTEREST ANALYSIS SUMMARY")

    def add_rule_11_60_analysis(self, rule_11_violations, rule_60_fraud):
        """Add Rule 11 sanctions and Rule 60(d)(3) fraud analysis sheets.

        Args:
            rule_11_violations: list of dicts with date, filer, filing,
                                misrepresentation, evidence_of_bad_faith,
                                applicable_rule, severity, remedy.
            rule_60_fraud: list of dicts with date, perpetrator, fraud_type,
                           description, evidence, legal_standard, severity,
                           procedural_impact.
        """
        if rule_11_violations:
            r11_cols = [
                "date", "filer", "filing", "misrepresentation",
                "evidence_of_bad_faith", "applicable_rule", "severity", "remedy"
            ]
            rows = []
            for v in rule_11_violations:
                row = {col: v.get(col, "") for col in r11_cols}
                for col in r11_cols:
                    if isinstance(row[col], list):
                        row[col] = "; ".join(str(x) for x in row[col])
                rows.append(row)

            df = pd.DataFrame(rows, columns=r11_cols)
            ws = self._add_sheet("Rule 11 Violations", df,
                                 "RULE 11 — SANCTIONS FOR MISREPRESENTATION")
            data_start_row = 4
            self._apply_autofilter(ws, 3, len(r11_cols))
            self._apply_severity_coloring(ws, df, "severity", data_start_row)

            for row_idx in range(data_start_row, data_start_row + len(df)):
                ws.cell(row=row_idx, column=1).number_format = "YYYY-MM-DD"

            ws.column_dimensions["C"].width = 35
            ws.column_dimensions["D"].width = 50
            ws.column_dimensions["E"].width = 45
            ws.column_dimensions["H"].width = 35
            ws.freeze_panes = "A4"

        if rule_60_fraud:
            r60_cols = [
                "date", "perpetrator", "fraud_type", "description",
                "evidence", "legal_standard", "severity", "procedural_impact"
            ]
            rows = []
            for f in rule_60_fraud:
                row = {col: f.get(col, "") for col in r60_cols}
                for col in r60_cols:
                    if isinstance(row[col], list):
                        row[col] = "; ".join(str(x) for x in row[col])
                rows.append(row)

            df = pd.DataFrame(rows, columns=r60_cols)
            ws = self._add_sheet("Rule 60 Fraud", df,
                                 "RULE 60(d)(3) — FRAUD ON THE COURT")
            data_start_row = 4
            self._apply_autofilter(ws, 3, len(r60_cols))
            self._apply_severity_coloring(ws, df, "severity", data_start_row)
            self._apply_row_severity_coloring(ws, df, "severity", data_start_row, len(r60_cols))

            for row_idx in range(data_start_row, data_start_row + len(df)):
                ws.cell(row=row_idx, column=1).number_format = "YYYY-MM-DD"

            ws.column_dimensions["D"].width = 55
            ws.column_dimensions["E"].width = 45
            ws.column_dimensions["F"].width = 40
            ws.column_dimensions["H"].width = 40
            ws.freeze_panes = "A4"

    def add_judicial_code_analysis(self, analysis):
        """Add per-judge judicial code compliance analysis sheets.

        Args:
            analysis: list of dicts, each with: judge_name,
                      overall_compliance_score, canons_violated (list of dicts
                      with canon_number, canon_title, violation_description,
                      date, evidence, severity), summary, recommendations.
        """
        if not analysis:
            return

        # Per-judge summary sheet
        summary_rows = []
        for judge in analysis:
            canons = judge.get("canons_violated", [])
            canon_nums = (", ".join(str(c.get("canon_number", "")) for c in canons)
                          if canons else "None")
            critical_count = sum(
                1 for c in canons if "CRITICAL" in str(c.get("severity", "")).upper()
            )
            high_count = sum(
                1 for c in canons if "HIGH" in str(c.get("severity", "")).upper()
            )
            summary_rows.append({
                "judge_name": judge.get("judge_name", ""),
                "overall_compliance_score": judge.get("overall_compliance_score", ""),
                "total_violations": len(canons),
                "critical_violations": critical_count,
                "high_violations": high_count,
                "canons_violated": canon_nums,
                "summary": judge.get("summary", ""),
                "recommendations": judge.get("recommendations", ""),
            })

        summary_df = pd.DataFrame(summary_rows)
        ws_summary = self._add_sheet("Judicial Code Summary", summary_df,
                                     "JUDICIAL CODE OF CONDUCT — COMPLIANCE SUMMARY")
        data_start_row = 4
        self._apply_autofilter(ws_summary, 3, len(summary_df.columns))

        # Conditional formatting on compliance score
        if "overall_compliance_score" in summary_df.columns:
            col_idx = list(summary_df.columns).index("overall_compliance_score") + 1
            col_letter = openpyxl.utils.get_column_letter(col_idx)
            last_row = data_start_row + len(summary_df) - 1
            cell_range = f"{col_letter}{data_start_row}:{col_letter}{last_row}"
            ws_summary.conditional_formatting.add(
                cell_range,
                CellIsRule(operator="lessThan", formula=["50"],
                           fill=self.ALERT_FILL)
            )
            ws_summary.conditional_formatting.add(
                cell_range,
                CellIsRule(operator="between", formula=["50", "70"],
                           fill=self.WARN_FILL)
            )
            ws_summary.conditional_formatting.add(
                cell_range,
                CellIsRule(operator="greaterThan", formula=["70"],
                           fill=self.GOOD_FILL)
            )

        ws_summary.column_dimensions["A"].width = 25
        ws_summary.column_dimensions["F"].width = 30
        ws_summary.column_dimensions["G"].width = 50
        ws_summary.column_dimensions["H"].width = 50
        ws_summary.freeze_panes = "A4"

        # Detailed canon violations sheet
        detail_rows = []
        for judge in analysis:
            for canon in judge.get("canons_violated", []):
                detail_rows.append({
                    "judge_name": judge.get("judge_name", ""),
                    "canon_number": canon.get("canon_number", ""),
                    "canon_title": canon.get("canon_title", ""),
                    "violation_description": canon.get("violation_description", ""),
                    "date": canon.get("date", ""),
                    "evidence": canon.get("evidence", ""),
                    "severity": canon.get("severity", ""),
                })

        if detail_rows:
            detail_df = pd.DataFrame(detail_rows)
            ws_detail = self._add_sheet("Judicial Canon Detail", detail_df,
                                        "JUDICIAL CANON VIOLATIONS — DETAIL")
            data_start_row = 4
            self._apply_autofilter(ws_detail, 3, len(detail_df.columns))
            self._apply_severity_coloring(ws_detail, detail_df, "severity", data_start_row)

            for row_idx in range(data_start_row, data_start_row + len(detail_df)):
                ws_detail.cell(row=row_idx, column=5).number_format = "YYYY-MM-DD"

            ws_detail.column_dimensions["C"].width = 30
            ws_detail.column_dimensions["D"].width = 55
            ws_detail.column_dimensions["F"].width = 50
            ws_detail.freeze_panes = "A4"

    def add_per_person_scoring(self, scoring):
        """Add per-person abuse scoring breakdown sheet.

        Args:
            scoring: list of dicts with person_name, role, total_score,
                     category_scores (dict mapping category to score),
                     incident_count, severity_distribution (dict),
                     key_findings (list of strings).
        """
        if not scoring:
            return

        # Collect all abuse categories across all persons
        all_categories = set()
        for person in scoring:
            for cat in person.get("category_scores", {}).keys():
                all_categories.add(cat)
        all_categories = sorted(all_categories)

        rows = []
        for person in scoring:
            row = {
                "person_name": person.get("person_name", ""),
                "role": person.get("role", ""),
                "total_score": person.get("total_score", 0),
                "incident_count": person.get("incident_count", 0),
            }
            cat_scores = person.get("category_scores", {})
            for cat in all_categories:
                row[f"score_{cat}"] = cat_scores.get(cat, 0)

            sev_dist = person.get("severity_distribution", {})
            row["critical_incidents"] = sev_dist.get("CRITICAL", 0)
            row["high_incidents"] = sev_dist.get("HIGH", 0)
            row["moderate_incidents"] = sev_dist.get("MODERATE", 0)
            row["low_incidents"] = sev_dist.get("LOW", 0)

            findings = person.get("key_findings", [])
            row["key_findings"] = "; ".join(str(f) for f in findings) if findings else ""
            rows.append(row)

        df = pd.DataFrame(rows)
        ws = self._add_sheet("Per-Person Scoring", df,
                             "PER-PERSON ABUSE SCORING BREAKDOWN")
        data_start_row = 4
        num_cols = len(df.columns)
        self._apply_autofilter(ws, 3, num_cols)

        # Color-scale on total_score column
        if "total_score" in df.columns:
            col_idx = list(df.columns).index("total_score") + 1
            col_letter = openpyxl.utils.get_column_letter(col_idx)
            last_row = data_start_row + len(df) - 1
            cell_range = f"{col_letter}{data_start_row}:{col_letter}{last_row}"
            ws.conditional_formatting.add(
                cell_range,
                ColorScaleRule(
                    start_type="min", start_color="C6EFCE",
                    mid_type="percentile", mid_value=50, mid_color="FFEB9C",
                    end_type="max", end_color="FFC7CE"
                )
            )

        # Data bars on per-category score columns
        for cat in all_categories:
            col_name = f"score_{cat}"
            if col_name in df.columns:
                col_idx = list(df.columns).index(col_name) + 1
                col_letter = openpyxl.utils.get_column_letter(col_idx)
                last_row = data_start_row + len(df) - 1
                cell_range = f"{col_letter}{data_start_row}:{col_letter}{last_row}"
                ws.conditional_formatting.add(
                    cell_range,
                    DataBarRule(start_type="min", end_type="max", color="FF6B6B")
                )

        # Widen key_findings column
        if "key_findings" in df.columns:
            findings_col_idx = list(df.columns).index("key_findings") + 1
            findings_letter = openpyxl.utils.get_column_letter(findings_col_idx)
            ws.column_dimensions[findings_letter].width = 60

        ws.freeze_panes = "A4"

    def add_dismissal_opportunities(self, opportunities):
        """Add sheet listing every point where the case should have been dismissed.

        Args:
            opportunities: list of dicts with date, event, legal_basis,
                           applicable_statute, why_not_dismissed,
                           responsible_official, prejudice_to_party,
                           severity, cumulative_impact.
        """
        if not opportunities:
            return

        columns = [
            "date", "event", "legal_basis", "applicable_statute",
            "why_not_dismissed", "responsible_official", "prejudice_to_party",
            "severity", "cumulative_impact"
        ]
        rows = []
        for opp in opportunities:
            row = {}
            for col in columns:
                val = opp.get(col, "")
                if isinstance(val, list):
                    val = "; ".join(str(v) for v in val)
                row[col] = val
            rows.append(row)

        df = pd.DataFrame(rows, columns=columns)

        try:
            df["date"] = pd.to_datetime(df["date"], errors="coerce")
            df = df.sort_values("date", na_position="last")
        except Exception:
            pass

        ws = self._add_sheet("Dismissal Opportunities", df,
                             "MISSED DISMISSAL OPPORTUNITIES")
        data_start_row = 4
        num_cols = len(columns)

        self._apply_autofilter(ws, 3, num_cols)
        self._apply_severity_coloring(ws, df, "severity", data_start_row)

        for row_idx in range(data_start_row, data_start_row + len(df)):
            ws.cell(row=row_idx, column=1).number_format = "YYYY-MM-DD"

        ws.column_dimensions["B"].width = 45
        ws.column_dimensions["C"].width = 45
        ws.column_dimensions["D"].width = 35
        ws.column_dimensions["E"].width = 45
        ws.column_dimensions["F"].width = 25
        ws.column_dimensions["G"].width = 40
        ws.column_dimensions["I"].width = 40

        # Summary row
        summary_row = data_start_row + len(df) + 1
        cell = ws.cell(row=summary_row, column=1, value="TOTAL MISSED OPPORTUNITIES")
        cell.font = Font(name="Calibri", size=11, bold=True)
        count_cell = ws.cell(row=summary_row, column=2, value=len(df))
        count_cell.font = Font(name="Calibri", size=11, bold=True)

        ws.freeze_panes = "A4"

    def add_judge_rotation_analysis(self, analysis):
        """Add judge rotation patterns analysis sheet.

        Args:
            analysis: dict with keys: rotations (list of dicts with judge_name,
                      start_date, end_date, days_served, reason_for_change,
                      orders_issued, significant_actions), pattern_analysis (dict),
                      anomalies (list of dicts).
        """
        if not analysis:
            return

        rotations = analysis.get("rotations", [])
        if rotations:
            rot_cols = [
                "judge_name", "start_date", "end_date", "days_served",
                "reason_for_change", "orders_issued", "significant_actions"
            ]
            rows = []
            for r in rotations:
                row = {col: r.get(col, "") for col in rot_cols}
                if isinstance(row.get("significant_actions"), list):
                    row["significant_actions"] = "; ".join(
                        str(a) for a in row["significant_actions"]
                    )
                rows.append(row)

            df = pd.DataFrame(rows, columns=rot_cols)
            ws = self._add_sheet("Judge Rotation", df,
                                 "JUDGE ROTATION ANALYSIS")
            data_start_row = 4
            self._apply_autofilter(ws, 3, len(rot_cols))

            for col_name in ["start_date", "end_date"]:
                col_idx = list(df.columns).index(col_name) + 1
                for row_idx in range(data_start_row, data_start_row + len(df)):
                    ws.cell(row=row_idx, column=col_idx).number_format = "YYYY-MM-DD"

            if "days_served" in df.columns:
                col_idx = list(df.columns).index("days_served") + 1
                for row_idx in range(data_start_row, data_start_row + len(df)):
                    ws.cell(row=row_idx, column=col_idx).number_format = "0"

            ws.column_dimensions["E"].width = 35
            ws.column_dimensions["G"].width = 55
            ws.freeze_panes = "A4"

        # Pattern analysis sub-sheet
        pattern = analysis.get("pattern_analysis", {})
        if pattern:
            self._add_dict_sheet("Rotation Patterns", pattern,
                                 "JUDGE ROTATION PATTERN ANALYSIS")

        # Anomalies sub-sheet
        anomalies = analysis.get("anomalies", [])
        if anomalies:
            anom_df = pd.DataFrame(anomalies)
            ws_anom = self._add_sheet("Rotation Anomalies", anom_df,
                                      "JUDGE ROTATION ANOMALIES")
            self._apply_autofilter(ws_anom, 3, len(anom_df.columns))
            self._apply_severity_coloring(ws_anom, anom_df, "severity", 4)
            ws_anom.freeze_panes = "A4"

    def add_conflict_of_interest(self, conflicts):
        """Add detected conflicts of interest sheet.

        Args:
            conflicts: list of dicts with official_name, role, conflict_type,
                       conflicting_party, description, evidence,
                       financial_interest, recusal_required, severity,
                       date_identified.
        """
        if not conflicts:
            return

        columns = [
            "official_name", "role", "conflict_type", "conflicting_party",
            "description", "evidence", "financial_interest", "recusal_required",
            "severity", "date_identified"
        ]
        rows = []
        for c in conflicts:
            row = {}
            for col in columns:
                val = c.get(col, "")
                if isinstance(val, list):
                    val = "; ".join(str(v) for v in val)
                elif isinstance(val, bool):
                    val = "YES" if val else "NO"
                row[col] = val
            rows.append(row)

        df = pd.DataFrame(rows, columns=columns)
        ws = self._add_sheet("Conflicts of Interest", df,
                             "DETECTED CONFLICTS OF INTEREST")
        data_start_row = 4
        num_cols = len(columns)

        self._apply_autofilter(ws, 3, num_cols)
        self._apply_severity_coloring(ws, df, "severity", data_start_row)

        # Highlight recusal_required = YES
        if "recusal_required" in df.columns:
            col_idx = list(df.columns).index("recusal_required") + 1
            for row_idx in range(data_start_row, data_start_row + len(df)):
                cell = ws.cell(row=row_idx, column=col_idx)
                if str(cell.value).upper() == "YES":
                    cell.fill = self.ALERT_FILL
                    cell.font = Font(name="Calibri", size=11, bold=True, color="C00000")

        # Format financial_interest as currency
        if "financial_interest" in df.columns:
            col_idx = list(df.columns).index("financial_interest") + 1
            for row_idx in range(data_start_row, data_start_row + len(df)):
                cell = ws.cell(row=row_idx, column=col_idx)
                try:
                    val = float(str(cell.value).replace("$", "").replace(",", ""))
                    cell.value = val
                    cell.number_format = "$#,##0.00"
                except (ValueError, TypeError):
                    pass

        # Format dates
        if "date_identified" in df.columns:
            col_idx = list(df.columns).index("date_identified") + 1
            for row_idx in range(data_start_row, data_start_row + len(df)):
                ws.cell(row=row_idx, column=col_idx).number_format = "YYYY-MM-DD"

        ws.column_dimensions["E"].width = 50
        ws.column_dimensions["F"].width = 45
        ws.freeze_panes = "A4"

    def add_charts_sheet(self, chart_paths):
        """Add a sheet with embedded chart images."""
        ws = self.wb.create_sheet(title="Charts")
        ws.merge_cells("A1:F1")
        cell = ws.cell(row=1, column=1, value="VISUAL ANALYSIS")
        cell.font = Font(name="Calibri", size=14, bold=True, color="1F4E79")

        row = 3
        for path in chart_paths:
            if os.path.exists(path):
                try:
                    img = openpyxl.drawing.image.Image(path)
                    img.width = 800
                    img.height = 400
                    ws.add_image(img, f"A{row}")
                    row += 22  # Space for next image
                except Exception:
                    ws.cell(row=row, column=1, value=f"[Chart: {os.path.basename(path)}]")
                    row += 2

    def add_raw_data_sheet(self, sheet_name, df):
        """Add a raw data sheet with auto-filters for querying in Excel."""
        if df is None or df.empty:
            return

        # Truncate sheet name to 31 chars (Excel limit)
        safe_name = sheet_name[:31]
        ws = self.wb.create_sheet(title=safe_name)

        # Write headers
        for col_idx, col_name in enumerate(df.columns, 1):
            cell = ws.cell(row=1, column=col_idx, value=str(col_name))
            cell.font = self.HEADER_FONT
            cell.fill = self.HEADER_FILL
            cell.alignment = Alignment(horizontal="center", wrap_text=True)

        # Write data
        for row_idx, row in enumerate(df.itertuples(index=False), 2):
            for col_idx, value in enumerate(row, 1):
                cell = ws.cell(row=row_idx, column=col_idx)
                if pd.isna(value):
                    cell.value = ""
                else:
                    cell.value = str(value) if not isinstance(value, (int, float)) else value

        # Auto-filter on all columns
        if len(df) > 0:
            ws.auto_filter.ref = ws.dimensions

        # Auto-width columns
        for col_idx, col_name in enumerate(df.columns, 1):
            max_len = max(len(str(col_name)), 10)
            for row in ws.iter_rows(min_row=2, max_row=min(len(df) + 1, 50),
                                     min_col=col_idx, max_col=col_idx):
                for cell in row:
                    if cell.value:
                        max_len = max(max_len, min(len(str(cell.value)), 50))
            ws.column_dimensions[openpyxl.utils.get_column_letter(col_idx)].width = max_len + 2

        # Freeze top row
        ws.freeze_panes = "A2"

    def save(self):
        """Save the workbook."""
        self.wb.save(self.output_path)
        return self.output_path


class WordReportGenerator:
    """Generates a formal Word document report."""

    def __init__(self, output_path="output/analysis_report.docx"):
        self.output_path = output_path
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Configure document styles."""
        style = self.doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    def add_title_page(self, case_name="", case_number="", state="", county=""):
        """Add a formal title page."""
        for _ in range(6):
            self.doc.add_paragraph()

        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("FORENSIC ANALYSIS REPORT")
        run.font.size = Pt(24)
        run.bold = True

        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("Post-Separation Abuse Pattern Analysis\n& Court Case Review")
        run.font.size = Pt(16)

        self.doc.add_paragraph()

        if case_name:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Case: {case_name}")
            run.font.size = Pt(14)

        if case_number:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(f"Case Number: {case_number}").font.size = Pt(12)

        if state or county:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            location = ", ".join(filter(None, [county, state]))
            p.add_run(f"Jurisdiction: {location}").font.size = Pt(12)

        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"Report Date: {datetime.now().strftime('%B %d, %Y')}").font.size = Pt(12)

        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("CONFIDENTIAL — ATTORNEY-CLIENT PRIVILEGED")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(192, 0, 0)

        self.doc.add_page_break()

    def add_table_of_contents(self):
        """Add table of contents placeholder."""
        self.doc.add_heading("Table of Contents", level=1)
        p = self.doc.add_paragraph(
            "[This table of contents should be updated after final report generation. "
            "Right-click and select 'Update Field' in Microsoft Word.]"
        )
        p.italic = True
        self.doc.add_page_break()

    def add_executive_summary(self, summary_data):
        """Add executive summary section."""
        self.doc.add_heading("1. Executive Summary", level=1)

        if isinstance(summary_data, dict):
            for key, value in summary_data.items():
                if isinstance(value, dict):
                    self.doc.add_heading(key.replace("_", " ").title(), level=3)
                    for k2, v2 in value.items():
                        p = self.doc.add_paragraph()
                        run = p.add_run(f"{k2.replace('_', ' ').title()}: ")
                        run.bold = True
                        p.add_run(str(v2))
                elif isinstance(value, list):
                    self.doc.add_heading(key.replace("_", " ").title(), level=3)
                    for item in value:
                        self.doc.add_paragraph(str(item), style="List Bullet")
                else:
                    p = self.doc.add_paragraph()
                    run = p.add_run(f"{key.replace('_', ' ').title()}: ")
                    run.bold = True
                    p.add_run(str(value))
        else:
            self.doc.add_paragraph(str(summary_data))

    def add_abuse_analysis_section(self, classified_df, formula_doc=None):
        """Add abuse analysis section."""
        self.doc.add_heading("2. Post-Separation Abuse Analysis", level=1)

        self.doc.add_heading("2.1 Classification Methodology", level=2)
        self.doc.add_paragraph(
            "Incidents were classified using a standardized taxonomy of post-separation "
            "abuse categories, each with defined behavioral traits and severity weights. "
            "The classification system is aligned with definitions from the Violence Against "
            "Women Act (VAWA) and established domestic violence research."
        )

        if formula_doc:
            self.doc.add_heading("2.2 Composite Abuse Severity Index (CASI)", level=2)
            self.doc.add_paragraph(f"Formula: {formula_doc.get('formula', 'N/A')}")
            components = formula_doc.get("components", {})
            for comp_id, comp_info in components.items():
                p = self.doc.add_paragraph()
                run = p.add_run(f"{comp_id} — {comp_info.get('name', '')}: ")
                run.bold = True
                p.add_run(comp_info.get("description", ""))

        if classified_df is not None and not classified_df.empty:
            self.doc.add_heading("2.3 Incident Classification Results", level=2)
            self._add_dataframe_table(classified_df.head(50))

    def add_cycle_prediction_section(self, predictions_data):
        """Add cycle prediction section."""
        self.doc.add_heading("3. Abuse Cycle Prediction", level=1)

        if "model_info" in predictions_data:
            info = predictions_data["model_info"]
            self.doc.add_paragraph(f"Data points analyzed: {info.get('data_points', 'N/A')}")
            self.doc.add_paragraph(f"Date range: {info.get('date_range', 'N/A')}")
            self.doc.add_paragraph(f"Method: {info.get('method', 'N/A')}")

        if "predictions" in predictions_data:
            self.doc.add_heading("3.1 Predicted Cycles", level=2)
            pred_df = pd.DataFrame(predictions_data["predictions"])
            self._add_dataframe_table(pred_df)

        # Disclaimer
        p = self.doc.add_paragraph()
        run = p.add_run("Disclaimer: ")
        run.bold = True
        run.font.color.rgb = RGBColor(192, 0, 0)
        p.add_run(predictions_data.get("model_info", {}).get("disclaimer",
                  "Predictions are statistical projections only."))

    def add_court_analysis_section(self, duration_data, filing_patterns, timeline):
        """Add court case analysis section."""
        self.doc.add_heading("4. Court Case Analysis", level=1)

        if duration_data and "error" not in duration_data:
            self.doc.add_heading("4.1 Case Duration", level=2)
            for key, value in duration_data.items():
                p = self.doc.add_paragraph()
                run = p.add_run(f"{key.replace('_', ' ').title()}: ")
                run.bold = True
                p.add_run(str(value))

        if filing_patterns:
            self.doc.add_heading("4.2 Filing Patterns", level=2)
            for key, value in filing_patterns.items():
                p = self.doc.add_paragraph()
                run = p.add_run(f"{key.replace('_', ' ').title()}: ")
                run.bold = True
                if isinstance(value, dict):
                    p.add_run(str(value))
                else:
                    p.add_run(str(value))

    def add_officials_section(self, officials_df, earnings, scorecards_df):
        """Add court officials analysis section."""
        self.doc.add_heading("5. Court Officials Analysis", level=1)

        if officials_df is not None and not officials_df.empty:
            self.doc.add_heading("5.1 Officials Involved", level=2)
            self._add_dataframe_table(officials_df)

        if earnings:
            self.doc.add_heading("5.2 Estimated Earnings from Case", level=2)
            earn_df = pd.DataFrame.from_dict(earnings, orient="index").reset_index()
            earn_df = earn_df.rename(columns={"index": "name"})
            self._add_dataframe_table(earn_df)

        if scorecards_df is not None and not scorecards_df.empty:
            self.doc.add_heading("5.3 Official Scorecards", level=2)
            for _, row in scorecards_df.iterrows():
                self.doc.add_heading(
                    f"{row.get('name', 'Unknown')} — {row.get('role', '')} "
                    f"(Grade: {row.get('overall_grade', 'N/A')})", level=3
                )
                self.doc.add_paragraph(str(row.get("summary", "")))
                if row.get("flags"):
                    self.doc.add_paragraph("Flags:", style="List Bullet")
                    for flag in str(row.get("flags", "")).split("; "):
                        if flag.strip():
                            self.doc.add_paragraph(flag.strip(), style="List Bullet 2")

    def add_violations_section(self, violations):
        """Add procedural violations section."""
        self.doc.add_heading("6. Procedural Violations", level=1)

        if not violations:
            self.doc.add_paragraph("No procedural violations detected in the available data.")
            return

        self.doc.add_paragraph(f"Total violations identified: {len(violations)}")

        for i, v in enumerate(violations, 1):
            self.doc.add_heading(
                f"Violation #{i}: {v.get('type', 'Unknown')}", level=2
            )
            p = self.doc.add_paragraph()
            run = p.add_run("Description: ")
            run.bold = True
            p.add_run(str(v.get("description", "")))

            p = self.doc.add_paragraph()
            run = p.add_run("Applicable Law: ")
            run.bold = True
            p.add_run(str(v.get("law", "")))

            p = self.doc.add_paragraph()
            run = p.add_run("Citation: ")
            run.bold = True
            p.add_run(str(v.get("citation", "")))

            p = self.doc.add_paragraph()
            run = p.add_run("Severity: ")
            run.bold = True
            severity = str(v.get("severity", ""))
            sev_run = p.add_run(severity)
            if "HIGH" in severity or "CRITICAL" in severity:
                sev_run.font.color.rgb = RGBColor(192, 0, 0)
                sev_run.bold = True

    def add_financial_section(self, costs_data):
        """Add financial analysis section."""
        self.doc.add_heading("7. Financial Analysis", level=1)

        if costs_data.get("total_actual"):
            self.doc.add_paragraph(f"Total documented costs: ${costs_data['total_actual']:,.2f}")
        if costs_data.get("total_estimated"):
            self.doc.add_paragraph(f"Total estimated costs: ${costs_data['total_estimated']:,.0f}")

        if costs_data.get("per_party"):
            self.doc.add_heading("7.1 Costs by Party", level=2)
            for party, amount in costs_data["per_party"].items():
                self.doc.add_paragraph(f"{party}: ${amount:,.2f}", style="List Bullet")

        if costs_data.get("per_official"):
            self.doc.add_heading("7.2 Payments to Officials", level=2)
            for official, amount in costs_data["per_official"].items():
                self.doc.add_paragraph(f"{official}: ${amount:,.2f}", style="List Bullet")

    def add_connections_section(self, connections):
        """Add connections analysis section."""
        self.doc.add_heading("8. Official Connections Analysis", level=1)

        if not connections:
            self.doc.add_paragraph("No connections identified from available data.")
            return

        for c in connections:
            p = self.doc.add_paragraph()
            run = p.add_run(f"{c['person_1']} ({c['role_1']}) ")
            run.bold = True
            p.add_run(f"↔ ")
            run = p.add_run(f"{c['person_2']} ({c['role_2']})")
            run.bold = True
            self.doc.add_paragraph(
                f"  Connection: {c['connection_type']} — {c.get('details', '')}",
                style="List Bullet"
            )

    def add_outcome_prediction_section(self, prediction_data):
        """Add case outcome prediction section."""
        self.doc.add_heading("9. Case Outcome Prediction", level=1)

        if isinstance(prediction_data, dict):
            for key, value in prediction_data.items():
                if key == "disclaimer":
                    p = self.doc.add_paragraph()
                    run = p.add_run("Disclaimer: ")
                    run.bold = True
                    p.add_run(str(value))
                elif isinstance(value, dict):
                    self.doc.add_heading(key.replace("_", " ").title(), level=2)
                    for k2, v2 in value.items():
                        self.doc.add_paragraph(
                            f"{k2.replace('_', ' ').title()}: {v2}",
                            style="List Bullet"
                        )
                else:
                    p = self.doc.add_paragraph()
                    run = p.add_run(f"{key.replace('_', ' ').title()}: ")
                    run.bold = True
                    p.add_run(str(value))

    def add_document_comparison_section(self, comparison_result):
        """Add document comparison section."""
        self.doc.add_heading("10. Document Comparison Analysis", level=1)

        self.doc.add_paragraph(
            f"Similarity Score: {comparison_result.get('similarity_score', 'N/A')}"
        )

        contras = comparison_result.get("contradictions", [])
        if contras:
            self.doc.add_heading("10.1 Contradictions Found", level=2)
            for c in contras:
                self.doc.add_paragraph(
                    f"Type: {c.get('type', '')} — "
                    f"Doc A: \"{c.get('doc_a_statement', c.get('doc_a_value', ''))[:100]}\" vs "
                    f"Doc B: \"{c.get('doc_b_statement', c.get('doc_b_value', ''))[:100]}\"",
                    style="List Bullet"
                )

    def add_charts(self, chart_paths):
        """Embed chart images in the document."""
        self.doc.add_heading("Appendix A: Visual Analysis", level=1)

        for path in chart_paths:
            if os.path.exists(path):
                try:
                    self.doc.add_picture(path, width=Inches(6))
                    self.doc.add_paragraph()
                except Exception:
                    self.doc.add_paragraph(f"[Chart: {os.path.basename(path)}]")

    def add_methodology_section(self):
        """Add methodology section for legal defensibility."""
        self.doc.add_heading("Appendix B: Methodology", level=1)

        self.doc.add_heading("Data Collection", level=2)
        self.doc.add_paragraph(
            "All data analyzed in this report was sourced from court records, "
            "documented communications, official filings, and evidence provided "
            "by the submitting party. No data was fabricated or assumed."
        )

        self.doc.add_heading("Classification Methodology", level=2)
        self.doc.add_paragraph(
            "Abuse incidents were classified using a standardized taxonomy of "
            "post-separation abuse categories. Each category includes defined "
            "behavioral traits with pre-assigned severity weights. Classification "
            "is performed by matching documented behaviors against trait definitions."
        )

        self.doc.add_heading("Statistical Methods", level=2)
        methods = [
            "Time-series analysis for temporal pattern identification",
            "Pearson correlation testing for variable relationships",
            "Linear regression for trend analysis and escalation detection",
            "Peak detection algorithms for cycle identification",
            "ANOVA testing for categorical variable significance",
        ]
        for m in methods:
            self.doc.add_paragraph(m, style="List Bullet")

        self.doc.add_heading("Limitations", level=2)
        self.doc.add_paragraph(
            "This analysis is limited by the completeness and accuracy of the "
            "source data. Statistical predictions are probabilistic in nature and "
            "should not be treated as certainties. This report does not constitute "
            "legal advice and should be reviewed by qualified professionals."
        )

        self.doc.add_heading("Reproducibility", level=2)
        self.doc.add_paragraph(
            "All calculations in this report are deterministic and fully reproducible "
            "from the source data using the documented methodology. The analysis "
            "software version and parameters are recorded for verification purposes."
        )

    def add_detailed_timeline_section(self, timeline):
        """Add detailed event timeline section to the Word report.

        Args:
            timeline: list of dicts with date, description, persons_involved,
                      laws_broken, human_rights_violations, method,
                      source_document, severity.
        """
        self.doc.add_heading("Appendix C: Detailed Event Timeline", level=1)

        if not timeline:
            self.doc.add_paragraph("No timeline events available.")
            return

        self.doc.add_paragraph(f"Total events documented: {len(timeline)}")
        self.doc.add_paragraph()

        # Group events by severity for narrative overview
        severity_groups = {}
        for event in timeline:
            sev = str(event.get("severity", "UNKNOWN")).upper()
            severity_groups.setdefault(sev, []).append(event)

        for sev in ["CRITICAL", "HIGH", "MODERATE", "LOW"]:
            events = severity_groups.get(sev, [])
            if not events:
                continue
            self.doc.add_heading(f"{sev} Severity Events ({len(events)})", level=2)
            for event in events:
                p = self.doc.add_paragraph()
                date_str = str(event.get("date", "Unknown date"))
                run = p.add_run(f"{date_str}: ")
                run.bold = True
                if sev == "CRITICAL":
                    run.font.color.rgb = RGBColor(192, 0, 0)
                p.add_run(str(event.get("description", "")))

                details = []
                persons = event.get("persons_involved", "")
                if persons:
                    if isinstance(persons, list):
                        persons = ", ".join(persons)
                    details.append(f"Persons: {persons}")
                laws = event.get("laws_broken", "")
                if laws:
                    if isinstance(laws, list):
                        laws = ", ".join(laws)
                    details.append(f"Laws broken: {laws}")
                hr_violations = event.get("human_rights_violations", "")
                if hr_violations:
                    if isinstance(hr_violations, list):
                        hr_violations = ", ".join(hr_violations)
                    details.append(f"Human rights violations: {hr_violations}")
                source = event.get("source_document", "")
                if source:
                    details.append(f"Source: {source}")

                for detail in details:
                    self.doc.add_paragraph(detail, style="List Bullet")

        # Full timeline table
        self.doc.add_heading("Complete Timeline Table", level=2)
        columns = ["date", "description", "persons_involved", "laws_broken",
                    "severity", "source_document"]
        rows = []
        for event in timeline:
            row = {}
            for col in columns:
                val = event.get(col, "")
                if isinstance(val, list):
                    val = "; ".join(str(v) for v in val)
                row[col] = str(val)[:200]
            rows.append(row)
        df = pd.DataFrame(rows, columns=columns)
        self._add_dataframe_table(df, max_rows=100)

    def add_legal_analysis_section(self, rule_11, rule_60, judicial_code):
        """Add legal analysis section covering Rule 11, Rule 60, and judicial code.

        Args:
            rule_11: list of dicts (Rule 11 violations).
            rule_60: list of dicts (Rule 60(d)(3) fraud findings).
            judicial_code: list of dicts (per-judge judicial code analysis).
        """
        self.doc.add_heading("Appendix D: Legal Analysis", level=1)

        # Rule 11
        self.doc.add_heading("D.1 Rule 11 — Sanctions for Misrepresentation", level=2)
        if not rule_11:
            self.doc.add_paragraph("No Rule 11 violations identified.")
        else:
            self.doc.add_paragraph(
                f"Total Rule 11 violations identified: {len(rule_11)}"
            )
            for i, v in enumerate(rule_11, 1):
                self.doc.add_heading(
                    f"Violation #{i}: {v.get('filing', 'Unknown Filing')}", level=3
                )
                p = self.doc.add_paragraph()
                run = p.add_run("Date: ")
                run.bold = True
                p.add_run(str(v.get("date", "N/A")))

                p = self.doc.add_paragraph()
                run = p.add_run("Filer: ")
                run.bold = True
                p.add_run(str(v.get("filer", "N/A")))

                p = self.doc.add_paragraph()
                run = p.add_run("Misrepresentation: ")
                run.bold = True
                p.add_run(str(v.get("misrepresentation", "")))

                p = self.doc.add_paragraph()
                run = p.add_run("Evidence of Bad Faith: ")
                run.bold = True
                evidence = v.get("evidence_of_bad_faith", "")
                if isinstance(evidence, list):
                    evidence = "; ".join(str(e) for e in evidence)
                p.add_run(str(evidence))

                p = self.doc.add_paragraph()
                run = p.add_run("Severity: ")
                run.bold = True
                severity = str(v.get("severity", ""))
                sev_run = p.add_run(severity)
                if "CRITICAL" in severity.upper() or "HIGH" in severity.upper():
                    sev_run.font.color.rgb = RGBColor(192, 0, 0)
                    sev_run.bold = True

                p = self.doc.add_paragraph()
                run = p.add_run("Recommended Remedy: ")
                run.bold = True
                p.add_run(str(v.get("remedy", "")))

        # Rule 60(d)(3)
        self.doc.add_heading("D.2 Rule 60(d)(3) — Fraud on the Court", level=2)
        if not rule_60:
            self.doc.add_paragraph("No Rule 60(d)(3) fraud findings identified.")
        else:
            self.doc.add_paragraph(
                f"Total fraud findings: {len(rule_60)}"
            )
            for i, f in enumerate(rule_60, 1):
                self.doc.add_heading(
                    f"Finding #{i}: {f.get('fraud_type', 'Unknown')}", level=3
                )
                for field in ["date", "perpetrator", "description", "evidence",
                              "legal_standard", "severity", "procedural_impact"]:
                    val = f.get(field, "")
                    if isinstance(val, list):
                        val = "; ".join(str(v) for v in val)
                    if val:
                        p = self.doc.add_paragraph()
                        run = p.add_run(f"{field.replace('_', ' ').title()}: ")
                        run.bold = True
                        sev_run = p.add_run(str(val))
                        if field == "severity" and (
                            "CRITICAL" in str(val).upper() or "HIGH" in str(val).upper()
                        ):
                            sev_run.font.color.rgb = RGBColor(192, 0, 0)
                            sev_run.bold = True

        # Judicial Code
        self.doc.add_heading("D.3 Judicial Code of Conduct Compliance", level=2)
        if not judicial_code:
            self.doc.add_paragraph("No judicial code analysis available.")
        else:
            for judge in judicial_code:
                name = judge.get("judge_name", "Unknown")
                score = judge.get("overall_compliance_score", "N/A")
                self.doc.add_heading(
                    f"Judge {name} — Compliance Score: {score}", level=3
                )

                summary_text = judge.get("summary", "")
                if summary_text:
                    self.doc.add_paragraph(str(summary_text))

                canons = judge.get("canons_violated", [])
                if canons:
                    self.doc.add_paragraph(
                        f"Canon violations ({len(canons)}):"
                    )
                    for canon in canons:
                        bullet_text = (
                            f"Canon {canon.get('canon_number', '?')}: "
                            f"{canon.get('canon_title', '')} — "
                            f"{canon.get('violation_description', '')}"
                        )
                        self.doc.add_paragraph(bullet_text, style="List Bullet")
                        if canon.get("evidence"):
                            self.doc.add_paragraph(
                                f"Evidence: {canon['evidence']}",
                                style="List Bullet 2"
                            )

                recs = judge.get("recommendations", "")
                if recs:
                    p = self.doc.add_paragraph()
                    run = p.add_run("Recommendations: ")
                    run.bold = True
                    p.add_run(str(recs))

    def add_jurisdiction_section(self, uccjea, dismissal_opps):
        """Add jurisdiction and dismissal analysis section.

        Args:
            uccjea: dict with UCCJEA analysis data.
            dismissal_opps: list of dicts with dismissal opportunity data.
        """
        self.doc.add_heading("Appendix E: Jurisdiction & Dismissal Analysis", level=1)

        # UCCJEA
        self.doc.add_heading("E.1 UCCJEA Jurisdiction Analysis", level=2)
        if not uccjea:
            self.doc.add_paragraph("No UCCJEA analysis available.")
        else:
            findings = uccjea.get("jurisdiction_findings", [])
            if findings:
                self.doc.add_heading("Jurisdiction Findings", level=3)
                for finding in findings:
                    p = self.doc.add_paragraph()
                    for key, value in finding.items():
                        if isinstance(value, list):
                            value = "; ".join(str(v) for v in value)
                        run = p.add_run(f"{key.replace('_', ' ').title()}: ")
                        run.bold = True
                        p.add_run(f"{value}  ")
                    self.doc.add_paragraph()

            addr_history = uccjea.get("address_history", [])
            if addr_history:
                self.doc.add_heading("Address History", level=3)
                addr_df = pd.DataFrame(addr_history)
                self._add_dataframe_table(addr_df)

            dismissal_grounds = uccjea.get("dismissal_grounds", {})
            if dismissal_grounds:
                self.doc.add_heading("UCCJEA Dismissal Grounds", level=3)
                for key, value in dismissal_grounds.items():
                    p = self.doc.add_paragraph()
                    run = p.add_run(f"{key.replace('_', ' ').title()}: ")
                    run.bold = True
                    if isinstance(value, list):
                        p.add_run("; ".join(str(v) for v in value))
                    elif isinstance(value, dict):
                        p.add_run(str(value))
                    else:
                        p.add_run(str(value))

        # Dismissal Opportunities
        self.doc.add_heading("E.2 Missed Dismissal Opportunities", level=2)
        if not dismissal_opps:
            self.doc.add_paragraph("No dismissal opportunities identified.")
        else:
            self.doc.add_paragraph(
                f"Total missed dismissal opportunities: {len(dismissal_opps)}"
            )
            for i, opp in enumerate(dismissal_opps, 1):
                date_str = str(opp.get("date", "Unknown"))
                event_str = str(opp.get("event", ""))
                self.doc.add_heading(
                    f"Opportunity #{i} — {date_str}", level=3
                )

                p = self.doc.add_paragraph()
                run = p.add_run("Event: ")
                run.bold = True
                p.add_run(event_str)

                for field in ["legal_basis", "applicable_statute",
                              "why_not_dismissed", "responsible_official",
                              "prejudice_to_party", "cumulative_impact"]:
                    val = opp.get(field, "")
                    if isinstance(val, list):
                        val = "; ".join(str(v) for v in val)
                    if val:
                        p = self.doc.add_paragraph()
                        run = p.add_run(f"{field.replace('_', ' ').title()}: ")
                        run.bold = True
                        p.add_run(str(val))

                severity = str(opp.get("severity", ""))
                if severity:
                    p = self.doc.add_paragraph()
                    run = p.add_run("Severity: ")
                    run.bold = True
                    sev_run = p.add_run(severity)
                    if "CRITICAL" in severity.upper() or "HIGH" in severity.upper():
                        sev_run.font.color.rgb = RGBColor(192, 0, 0)
                        sev_run.bold = True

    def add_best_interest_section(self, factors):
        """Add best interest of the child analysis section.

        Args:
            factors: dict with parent_a_name, parent_b_name, factors (list),
                     summary (dict).
        """
        self.doc.add_heading("Appendix F: Best Interest of the Child Analysis", level=1)

        if not factors:
            self.doc.add_paragraph("No best interest analysis available.")
            return

        parent_a = factors.get("parent_a_name", "Parent A")
        parent_b = factors.get("parent_b_name", "Parent B")
        factor_list = factors.get("factors", [])

        self.doc.add_paragraph(
            "The following analysis applies the statutory best-interest-of-the-child "
            "factors. Each factor is scored on a 1-10 scale for each parent, with "
            "supporting evidence references. Scores of 1-3 indicate significant "
            "concerns, 4-6 indicate moderate performance, and 7-10 indicate strong "
            "performance."
        )

        if factor_list:
            self.doc.add_heading("Factor-by-Factor Analysis", level=2)

            # Build summary table
            table_rows = []
            for f in factor_list:
                table_rows.append({
                    "Factor #": f.get("factor_number", ""),
                    "Factor": f.get("factor_name", ""),
                    f"{parent_a} Score": f.get("parent_a_score", ""),
                    f"{parent_b} Score": f.get("parent_b_score", ""),
                })
            summary_df = pd.DataFrame(table_rows)
            self._add_dataframe_table(summary_df)

            self.doc.add_paragraph()

            # Detailed narrative per factor
            for f in factor_list:
                factor_num = f.get("factor_number", "?")
                factor_name = f.get("factor_name", "")
                self.doc.add_heading(
                    f"Factor {factor_num}: {factor_name}", level=3
                )

                score_a = f.get("parent_a_score", "N/A")
                score_b = f.get("parent_b_score", "N/A")

                p = self.doc.add_paragraph()
                run = p.add_run(f"{parent_a} Score: ")
                run.bold = True
                score_run = p.add_run(str(score_a))
                try:
                    if int(score_a) <= 3:
                        score_run.font.color.rgb = RGBColor(192, 0, 0)
                        score_run.bold = True
                    elif int(score_a) >= 7:
                        score_run.font.color.rgb = RGBColor(0, 128, 0)
                        score_run.bold = True
                except (ValueError, TypeError):
                    pass

                evidence_a = f.get("parent_a_evidence", "")
                if evidence_a:
                    self.doc.add_paragraph(
                        f"Evidence: {evidence_a}", style="List Bullet"
                    )

                p = self.doc.add_paragraph()
                run = p.add_run(f"{parent_b} Score: ")
                run.bold = True
                score_run = p.add_run(str(score_b))
                try:
                    if int(score_b) <= 3:
                        score_run.font.color.rgb = RGBColor(192, 0, 0)
                        score_run.bold = True
                    elif int(score_b) >= 7:
                        score_run.font.color.rgb = RGBColor(0, 128, 0)
                        score_run.bold = True
                except (ValueError, TypeError):
                    pass

                evidence_b = f.get("parent_b_evidence", "")
                if evidence_b:
                    self.doc.add_paragraph(
                        f"Evidence: {evidence_b}", style="List Bullet"
                    )

        # Summary
        summary = factors.get("summary", {})
        if summary:
            self.doc.add_heading("Overall Assessment", level=2)
            for key, value in summary.items():
                p = self.doc.add_paragraph()
                run = p.add_run(f"{key.replace('_', ' ').title()}: ")
                run.bold = True
                if isinstance(value, list):
                    p.add_run("; ".join(str(v) for v in value))
                else:
                    p.add_run(str(value))

    def _add_dataframe_table(self, df, max_rows=50):
        """Add a DataFrame as a formatted Word table."""
        if df.empty:
            return

        df = df.head(max_rows)
        table = self.doc.add_table(rows=1, cols=len(df.columns), style="Light Shading Accent 1")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Headers
        for i, col in enumerate(df.columns):
            cell = table.rows[0].cells[i]
            cell.text = str(col).replace("_", " ").title()
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)

        # Data
        for _, row_data in df.iterrows():
            row = table.add_row()
            for i, value in enumerate(row_data):
                cell = row.cells[i]
                if pd.isna(value):
                    cell.text = ""
                else:
                    cell.text = str(value)[:200]
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    def save(self):
        """Save the Word document."""
        self.doc.save(self.output_path)
        return self.output_path
